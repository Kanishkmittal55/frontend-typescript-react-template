import re
import logging
import tiktoken
from utils.unstructured_data_utils import (
    nodesTextToListOfDict,
    relationshipTextToListOfDict,
)
import httpx
import psutil
from typing import Optional
import json
from driver.neo4j import Neo4jDatabase
from typing import Callable, List, Dict, Any
import os
# ─── near the other imports ──────────────────────────
from components.embeddings import SimpleEmbedder          # ← your helper class
from threading import Lock




openai_key = os.getenv("OPENAI_API_KEY")
if not openai_key:
    raise RuntimeError("Missing OPENAI_API_KEY in environment")


def generate_system_message() -> str:
    return """
You are a data scientist working for a company that is building a graph database. Your task is to extract information from data and convert it into a graph database.
Provide a set of Nodes in the form [ENTITY_ID, TYPE, PROPERTIES] and a set of relationships in the form [ENTITY_ID_1, RELATIONSHIP, ENTITY_ID_2, PROPERTIES].
It is important that the ENTITY_ID_1 and ENTITY_ID_2 exists as nodes with a matching ENTITY_ID. If you can't pair a relationship with a pair of nodes don't add it.
When you find a node or relationship you want to add try to create a generic TYPE for it that  describes the entity you can also think of it as a label.

Example:
Data: Alice lawyer and is 25 years old and Bob is her roommate since 2001. Bob works as a journalist. Alice owns a the webpage www.alice.com and Bob owns the webpage www.bob.com.
Nodes: ["alice", "Person", {"age": 25, "occupation": "lawyer", "name":"Alice"}], ["bob", "Person", {"occupation": "journalist", "name": "Bob"}], ["alice.com", "Webpage", {"url": "www.alice.com"}], ["bob.com", "Webpage", {"url": "www.bob.com"}]
Relationships: ["alice", "roommate", "bob", {"start": 2021}], ["alice", "owns", "alice.com", {}], ["bob", "owns", "bob.com", {}]
"""


def num_tokens_from_string(string: str) -> int: 
    """
    Estimate the number of tokens in a string using the LLaMA tokenizer.
    :param string: The input string.
    :return: The number of tokens in the string.
    """

    # Use tiktoken's cl100k_base tokenizer to encode the input string
    tokenizer = tiktoken.get_encoding("cl100k_base")
    tokens = tokenizer.encode(string)
    return len(tokens)

def max_allowed_token_length() -> int:
    """
    Return the maximum number of tokens the model can handle.
    :return: The maximum allowed token length.
    """
    return 1024


def splitString(string, max_length) -> List[str]:
    return [string[i : i + max_length] for i in range(0, len(string), max_length)]


def splitStringToFitTokenSpace(string: str, token_use_per_string: int) -> List[str]:
    allowed_tokens = max_allowed_token_length() - token_use_per_string
    chunked_data = splitString(string, 500)  # Split based on approximate length
    # print("The length of the splitted text array is: ", len(chunked_data))

    # print(chunked_data[0])

    combined_chunks = []
    current_chunk = ""
    
    for chunk in chunked_data:
        # Calculate token count for the combined chunk
        current_chunk_tokens = num_tokens_from_string(current_chunk)
        # print("Length of the current chunk of 500 words based on regex is in tokens is :", current_chunk_tokens)
        chunk_tokens = num_tokens_from_string(chunk)
        # print("Length of the chunk token of 500 words based on regex is in tokens is :", current_chunk_tokens)

        if current_chunk_tokens + chunk_tokens <= allowed_tokens:
            current_chunk += chunk
        else:
            # Append the non-empty current chunk and start a new one
            if current_chunk.strip():
                combined_chunks.append(current_chunk.strip())
            current_chunk = chunk
    
    # Append any remaining chunk after the loop
    if current_chunk.strip():
        combined_chunks.append(current_chunk.strip())

    return combined_chunks



async def old_process(chunk: str) -> str:
    """
    Process a single chunk by making an asynchronous request to the Ollama endpoint.
    :param chunk: The chunk to process.
    :return: The processed response as a string.
    """
    try:
        # Prepare the request payload with only the chunk (prompt)
        payload = {
            "prompt": chunk  # Assuming the API expects a field named "chunk" for the prompt
        }

        # Log the payload
        logging.debug(f"Sending request to http://localhost:7860/ollama/chat with payload: {payload}")

        # Use httpx for asynchronous requests
        async with httpx.AsyncClient() as client:
            response = await client.post("http://localhost:7860/ollama/chat", json=payload)

        # Raise an exception for HTTP errors
        response.raise_for_status()

        # Parse and return the generated response
        # Access the 'message' field and then 'content
        result = response.json()

        print("The response is :",result)        
        # From generated text
        message = result.get('generated_text', {}).get('message', {})

        # Extract the content
        content = message.get('content', None)

        # print(content)
        
        return content
    except httpx.RequestError as e:
        logging.error(f"Error communicating with Ollama: {e}")
        return f"Error: {e}"
    except ValueError as e:  # Handle JSON decoding errors
        logging.error(f"Invalid JSON received from Ollama: {response.text}")
        return f"Invalid JSON received: {response.text}"
        

def generate_prompt(data) -> str:
    return f"""
Data: {data}"""


def clean_llm_response(raw_response: str) -> str:
    """
    Cleans the LLM response to remove any notes or extra comments,
    and extracts only the Nodes and Relationships section.
    Handles bullet points and additional formatting issues.
    
    :param raw_response: The raw response from the LLM.
    :return: Cleaned string containing only nodes and relationships.
    """
    # Use regex to extract the **Nodes** and **Relationships** sections
    match = re.search(r"Nodes:\s*(.*?)Relationships:\s*(.*)", raw_response, re.S)
    if match:
        nodes = re.sub(r"^\s*-\s*", "", match.group(1).strip(), flags=re.M)  # Remove bullet points from Nodes
        relationships = re.sub(r"^\s*-\s*", "", match.group(2).strip(), flags=re.M)  # Remove bullet points from Relationships

        # Remove any trailing note or comments
        relationships = re.split(r"(?i)\n(?:Note|Let me know)", relationships, maxsplit=1)[0].strip()

        # Reconstruct cleaned response in the expected format
        cleaned_response = f"Nodes: {nodes}\nRelationships: {relationships}"
        return cleaned_response
    else:
        # Return empty structure if no match is found
        return "Nodes: []\nRelationships: []"




def getNodesAndRelationshipsFromResult(result):
    regex = "Nodes:\s+(.*?)\s?\s?Relationships:\s?\s?(.*)"
    internalRegex = "\[(.*?)\]"
    nodes = []
    relationships = []

    print("Starting to parse :" , result, len(result))
    
    for row in result:
        print(f"Processing row: {row}")
        
        # Match using regex
        parsing = re.match(regex, row, flags=re.S)
        if parsing is None:
            print(f"No match found for row: {row}")
            continue
        
        print("Regex match found. Extracting nodes and relationships.")
        
        # Extract raw nodes and relationships
        rawNodes = str(parsing.group(1))
        rawRelationships = parsing.group(2)
        
        print(f"Extracted rawNodes: {rawNodes}")
        print(f"Extracted rawRelationships: {rawRelationships}")
        
        # Find individual elements using internal regex
        nodes.extend(re.findall(internalRegex, rawNodes))
        relationships.extend(re.findall(internalRegex, rawRelationships))
        
        print(f"Current list of nodes: {nodes}")
        print(f"Current list of relationships: {relationships}")

    # Validate extracted nodes and relationships before inserting them
    nodes = [node for node in nodes if len(node.split(",")) >= 2]  # Ensure each node has at least name and label
    relationships = [rel for rel in relationships if len(rel.split(",")) >= 3]  # Ensure valid relationship structure
    
    if not nodes:
        print("Warning: No valid nodes found.")
    if not relationships:
        print("Warning: No valid relationships found.")

    # Convert raw text data to structured dictionaries
    print("Converting extracted nodes and relationships to structured dictionaries.")
    try:
        result = dict()
        result["nodes"] = nodesTextToListOfDict(nodes)
        result["relationships"] = relationshipTextToListOfDict(relationships)
    except Exception as e:
        logging.error(f"Error during conversion of nodes or relationships: {e}")
        raise
    
    print("Inside the function nodes are : ", result["nodes"])

    logging.info("Successfully parsed nodes and relationships.")
    return result


async def run_with_chunk_logging(data: str) -> List[str]:
        print("Process Started with the patent text")


        ### Part 1 - Splitting the chunk text.

        system_message = generate_system_message() # Currently 324 tokens are only used for the system message or the intruction of the task.
        prompt_string = generate_prompt(data)
        print("No. of tokens in the prompt input string :", num_tokens_from_string(prompt_string))
        print("No. of tokens in the system message string :", num_tokens_from_string(system_message) )
        
        token_usage_per_prompt = num_tokens_from_string(
            system_message + prompt_string
        )

        print("token usage per prompt :",  token_usage_per_prompt)

        # Once we get the total token to be used including the system prompt we split it to fit the token space.
        chunked_data = splitStringToFitTokenSpace(
            string=data, token_use_per_string=num_tokens_from_string(system_message)
        )



        ### Part 2 - Processing the chunk sequentially the new code we have is concurrent in nature.

        print("So number of chunks created from the text are : ",len(chunked_data))

        results = []
        labels = set()
        chunks = []  # To store chunk metadata
        # print ( "Chunks have been created")

        print("Starting chunkwise processing")
        
        chunkResult = {} # getNodesAndRelationshipsFromResult returns this -> dict[str, list[dict[str, Any]]]

        for i,chunk in enumerate(chunked_data, start=1):
            


            print(f"Chunk number {i} chunk sent: {chunk}")
            
            # Log memory usage before processing the chunk
            memory_info = psutil.virtual_memory()
            print(f"Memory before processing chunk {i}: {memory_info.used / (1024**2):.2f} MB")

            processedChunk = await old_process(chunk)
            print(f"Chunk number {i} processedChunk : ", processedChunk)
            cleaned_response = clean_llm_response(processedChunk)
            print("The cleaned response is : ", cleaned_response)
            chunkResult = getNodesAndRelationshipsFromResult([cleaned_response])
            print("chunkResult- nodes and relationships : ", chunkResult.get("nodes", []))
            newLabels = [node["label"] for node in chunkResult["nodes"]]
            print("newLabels", newLabels)
            results.append(cleaned_response) # Why cleaned reponse ? Speculate
            labels.update(newLabels)
            # Append chunk metadata and result to chunks array
            chunks.append({
                "chunk_number": i,
                "system_prompt": system_message,
                "input_chunk_text": chunk,
                "chunk_result_nodes": chunkResult.get("nodes", []),  
                # Ensure no KeyError ,by providing a default empty key [] 
                # The empty brakcets serve as a default value in case the key "nodes" is missing from the dictionary.
                "chunk_result_relationships": chunkResult.get("relationships", [])
            })

            # Log memory usage after processing the chunk
            memory_info = psutil.virtual_memory()
            print(f"Memory after processing chunk {i}: {memory_info.used / (1024**2):.2f} MB")

        ## So No matter how many time one runs this "getNodesAndRelationshipsFromResult" function
        ## it will lead to further refinement of results.
        final_result = getNodesAndRelationshipsFromResult(results)
        return final_result, chunks




#####################################################################################
# The product report and discovery workflow
#####################################################################################

from llm.openai import OpenAIChat


# Now retrieve the API key as a string
api_key = openai_key

# Initialize LLM 
llm = OpenAIChat(
    openai_api_key=api_key, model_name="gpt-4o-mini", max_tokens=4096
)

# Helper Functions - 
async def openai_generate(prompt: str) -> str:
    """
    Sends a prompt to the OpenAI endpoint and returns the raw output.

    :param prompt: The input prompt to send to the OpenAI API.
    :return: The raw output generated by the OpenAI API.
    """
    messages = [
        {"role": "user", "content": prompt}  # Construct the message payload
    ]
    print(f"Sending request to OpenAI endpoint with messages: {messages}")
    
    # Assuming `llm.generate()` is correctly configured to accept the messages
    output = llm.generate(messages)
    print("The output is:", output)
    return output

async def process(chunk: str, provider: str) -> str:
    """
    Process a single chunk by making an asynchronous request to the specified provider's endpoint.
    :param chunk: The chunk to process.
    :param provider: The provider name (e.g., "ollama", "openai", "groq").
    :return: The processed response as a string.
    """
    try:
        # Prepare the request payload with only the chunk (prompt)
        payload = {
            "prompt": chunk  # Assuming the API expects a field named "chunk" for the prompt
        }

        # Route to appropriate provider using a match-case structure
        if provider == "ollama":
            logging.debug(f"Sending request to http://localhost:7860/ollama/chat with payload: {payload}")
            async with httpx.AsyncClient() as client:
                response = await client.post("http://localhost:7860/ollama/chat", json=payload)
            response.raise_for_status()
            result = response.json()
            message = result.get('generated_text', {}).get('message', {})
            content = message.get('content', None)

        elif provider == "openai":
            output = await openai_generate(chunk)
            return output

        elif provider == "groq":
            logging.debug(f"Sending request to Groq endpoint with payload: {payload}")
            async with httpx.AsyncClient() as client:
                response = await client.post("http://localhost:7860/groq/chat", json=payload)
            response.raise_for_status()
            result = response.json()
            content = result.get('response', None)  # Adjust based on Groq's API response structure

        else:
            raise ValueError(f"Unsupported provider: {provider}")

        return content

    except httpx.RequestError as e:
        logging.error(f"Error communicating with {provider}: {e}")
        return f"Error: {e}"
    except ValueError as e:
        logging.error(f"Invalid JSON received from {provider}: {e}")
        return f"Invalid JSON received: {e}"

def cleaned_name_and_description(response):
    """
    Extracts the product name and description from the response.
    
    :param response: The raw response string from the OpenAI API.
    :return: A dictionary containing the product name and description.
    """
    import re

    # Regular expressions to capture the name and description
    name_pattern = r"\*\*Product Type Name:\*\* (.+)"
    description_pattern = r"\*\*Description:\*\* (.+)"

    # Extract name and description
    name_match = re.search(name_pattern, response)
    description_match = re.search(description_pattern, response)

    # Get the matched content or default to None
    product_name = name_match.group(1).strip() if name_match else None
    product_description = description_match.group(1).strip() if description_match else None

    # Return as a dictionary
    return {
        "name": product_name,
        "description": product_description
    }

async def extract_name_description(extracted_info: str, provider: str) -> Optional[Dict[str, str]]:
    """
    Process the complete extracted information to finalize product name and description based on the provider.
    Input:
        extracted_info: string of extracted information
        provider: provider name (e.g., 'openai', 'groq', 'ollama')
    Output:
        A final refined string with product name and description
    """
    try:
        # Format the prompt
        prompt = f"""
        ### Extracted Information:
        {extracted_info}

        ### Instruction:
        You are a data scientist working for a company that is building a report for a cosmetic patent document. You are an author and have the capability to provide a very appropriate type of patent products based on gathered description. Provide an appropriate product type name describing very aptly what the product is about in a few words and a 2-sentence description of the product based on the extracted information, which represents a summary of the patent document.
        """

        # Prepare the request payload
        payload = {
            "prompt": prompt
        }

        logging.debug(f"Prepared payload: {payload}")

        # Route to appropriate provider using a match-case structure
        content = None

        if provider == "ollama":
            logging.debug(f"Sending request to http://localhost:7860/ollama/chat with payload: {payload}")
            async with httpx.AsyncClient() as client:
                response = await client.post("http://localhost:7860/ollama/chat", json=payload)
            response.raise_for_status()
            result = response.json()
            message = result.get('generated_text', {}).get('message', {})
            content = message.get('content', None)

        elif provider == "openai":
            output = await openai_generate(prompt)
            cleaned_output = cleaned_name_and_description(output)
            return cleaned_output

        elif provider == "groq":
            logging.debug(f"Sending request to Groq endpoint with payload: {payload}")
            async with httpx.AsyncClient() as client:
                response = await client.post("http://localhost:7860/groq/chat", json=payload)
            response.raise_for_status()
            result = response.json()
            content = result.get('response', None)  # Adjust based on Groq's API response structure

        else:
            raise ValueError(f"Unsupported provider: {provider}")

        if content:
            return content
        else:
            logging.error("No content received in the response.")
            return None

    except httpx.RequestError as e:
        logging.error(f"Error communicating with {provider}: {e}")
        return f"Error: {e}"

    except ValueError as e:  # Handle JSON decoding errors and invalid provider
        logging.error(f"ValueError: {e}")
        return f"ValueError: {e}"

    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        return f"Unexpected error: {e}"


def clean_functional_role_info(received_text: str) -> Optional[Dict]:
    """
    Transforms the received functional role information text into a JSON object.
    
    :param received_text: The raw string containing functional role information.
    :return: A JSON object with functional roles, or None if parsing fails.
    """
    try:
        # Attempt to parse the received text as JSON
        functional_roles_json = json.loads(received_text)

        # Validate that the required "functional_roles" key exists
        if "functional_roles" in functional_roles_json:
            return functional_roles_json  # Return the JSON object if valid

    except json.JSONDecodeError as e:
        print(f"Failed to parse JSON: {e}")

    # Return None if parsing fails or required key is missing
    return None

async def final_composition_information(extracted_info: str, provider: str) -> Optional[Dict[str, str]]:
    """
    Process the complete extracted information to finalize product composition report based on the provider.
    Input:
        extracted_info: string of extracted information
        provider: provider name (e.g., 'openai', 'groq', 'ollama')
    Output:
        A JSON object with the product's functional roles, chemicals, and weights
    """
    output_format = """
     {
        "functional_roles": {
            "Emollient": [
            {"chemical": "Dimethicone", "weight": "15%"},
            {"chemical": "Caprylic Triglyceride", "weight": "10%"}
            ],
            "Humectant": [
            {"chemical": "Glycerin", "weight": "20%"}
            ],
            "Preservative": [
            {"chemical": "Phenoxyethanol", "weight": "1%"}
            ]
        }
        }
    """
    try:
        # Format the prompt
        prompt = f"""
        ### Extracted Information:
        {extracted_info}

        ### Instruction:
        You are a data scientist working for a company that is building a report for a cosmetic patent document. Your job is to provide the final composition report of the product mentioned in the patent. Using the provided extracted functional role information, compile a comprehensive JSON object detailing the product's functional roles, the chemicals belonging to those roles, and their weights in percentages or ranges.Include all the unique extracted functional roles, their chemicals, and weights (or weight ranges). Ensure the JSON is clean, well-structured, and free of duplicates.If no functional roles are found in the complete data, respond with an empty functional_roles JSON object. Dont add comments, notes, or suggestions. Only provide the JSON object in the specified format.

        ### Output Format:
        Provide the response strictly in the following JSON format:
        {output_format}
        """

        # Prepare the request payload
        payload = {
            "prompt": prompt
        }

        logging.debug(f"Prepared payload: {payload}")

        # Route to appropriate provider using a match-case structure
        content = None

        if provider == "ollama":
            logging.debug(f"Sending request to http://localhost:7860/ollama/chat with payload: {payload}")
            async with httpx.AsyncClient() as client:
                response = await client.post("http://localhost:7860/ollama/chat", json=payload)
            response.raise_for_status()
            result = response.json()
            message = result.get('generated_text', {}).get('message', {})
            content = message.get('content', None)

        elif provider == "openai":
            output = await openai_generate(prompt)
            cleaned_output = clean_functional_role_info(output)
            return cleaned_output

        elif provider == "groq":
            logging.debug(f"Sending request to Groq endpoint with payload: {payload}")
            async with httpx.AsyncClient() as client:
                response = await client.post("http://localhost:7860/groq/chat", json=payload)
            response.raise_for_status()
            result = response.json()
            content = result.get('response', None)  # Adjust based on Groq's API response structure

        else:
            raise ValueError(f"Unsupported provider: {provider}")

        if content:
            return content
        else:
            logging.error("No content received in the response.")
            return None

    except httpx.RequestError as e:
        logging.error(f"Error communicating with {provider}: {e}")
        return f"Error: {e}"

    except ValueError as e:  # Handle JSON decoding errors and invalid provider
        logging.error(f"ValueError: {e}")
        return f"ValueError: {e}"

    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        return f"Unexpected error: {e}"

# Function to clean and tokenize text into words
def clean_and_tokenize(text: str) -> set:
    """
    Cleans and tokenizes input text into a set of words.
    Removes punctuation and converts to lowercase.
    """
    return set(re.findall(r'\b\w+\b', text.lower()))

# Function to check and append only unique words
def information_checker(information_extracted: str, new_response: str) -> str:
    """
    Compares existing extracted information with new response, appends only unique words.
    """
    # Tokenize and clean existing and new responses
    existing_tokens = clean_and_tokenize(information_extracted)
    new_tokens = clean_and_tokenize(new_response)
    
    # Find unique words from new response
    unique_tokens = new_tokens - existing_tokens

    if unique_tokens:
        # Append unique words to the extracted information
        updated_information = f"{information_extracted.strip()} {' '.join(unique_tokens)}"
        print("New unique words added:", unique_tokens)
        return updated_information.strip()
    else:
        print("No new unique words to append.")
        return information_extracted.strip()

functional_roles = "Antioxidant, Humectant, Emollient, Surfactant, Emulsifier, Preservative, Fragrance, Colorant, UV Filter (Sunscreen Agent), Thickener/Viscosity Modifier, Conditioning Agent, Astringent, Film-Former, Opacifier, Solvent, Exfoliant, Antimicrobial Agent, Chelating Agent, Antifoaming Agent, Moisturizer, Absorbent, Mattifier, Skin Protectant, Soothing Agent, Exfoliating Enzyme, Wetting Agent, Texturizer, Anti-inflammatory, Desensitizer, Penetration Enhancer, Hair Fixative, Antidandruff Agent, Anti-aging Agent, Brightening Agent, Anti-acne Agent, Lubricant, Deodorant Agent, Toning Agent, Antiperspirant, Styling Agent, Hair Growth Stimulator, Anti-hair Loss Agent, Nail Hardener, Plasticizer, Peptide/Protein Agent, Anti-pollution Agent, Anti-oxidative Stress Agent"

def extract_document_details(text_input):
    """
    Helper function to extract patent document details from the input text
    and return the extracted JSON object.

    Args:
        text_input (str): Text containing the patent details in JSON-like format.

    Returns:
        dict: Extracted patent document details as a JSON object.
    """
    try:
        # Extract JSON object from the text input
        start_index = text_input.find("{")
        end_index = text_input.rfind("}") + 1
        json_text = text_input[start_index:end_index]
        document_details = json.loads(json_text)
        print("The final document details are :", document_details)
        return document_details.get("patent", {})

    except (ValueError, KeyError, AttributeError) as e:
        print(f"Error processing document details: {e}")
        return {}

async def product_discovery_workflow(data: str, provider: str) -> List[dict]:

    ### the user gave it a schema based on which a schema template is populated or it could be asked in the form of a react form

    ### once the schema has been converted to the prompt we can use that prompt to create a query centric graph and the use that graph 

    ### to query the LLM , the point was to reduce validate ground truth and 

    print("Process Started with the patent text")

    # Split data into chunks to fit token space
    max_tokens_per_chunk = 4096  # Total token budget
    chunked_data = splitStringToFitTokenSpace(string=data, token_use_per_string=0)  # Adjust chunk size
    print("Number of chunks created from the text:", len(chunked_data))
    total_token_count = 0  # Tracks the total tokens processed

    # results = []
    information_extracted = ""  # Accumulator
    composition_information_extracted = "" # Accumulator
    document_information_extracted = "" # Accumulator

    print("Starting chunkwise processing")

    for i, chunk in enumerate(chunked_data, start=1):
        print(f"\nProcessing Chunk {i}:")

        # Format prompt and injecting contextual data to improve answer quality to include previous context and the current chunk
        prompt = f"""
        ### Information Extracted Till Now (50 words max):
        {information_extracted}

        ### New Chunk:
        {chunk}

        ### Instruction:
        You are a data scientist working for a company that is building a report for a cosmetic patent document. Summarize the 'New Chunk' in no more than **30 keywords or tokens**, focusing only on product name, description. 
        If no relevant information is found, respond: "No new information found." Only add the words that would be used at
        the last chunk to create one holistic product name and description from the complete information. You will not add any suggestions , comments or notes to the response, just give the summarized paragraph containing keywords that would help in the final analysis.
        """

        composition_prompt = f"""
        ### Functional Role Extracted Till Now:
        {composition_information_extracted}

        ### New Chunk:
        {chunk}

        ### Instruction:
        You are a data scientist working for a company that is building a report for a cosmetic patent document. Your job is to find the chemicals used to make the product and in what percentage weights or weight ranges they have been used in the product formulation. Your task is to extract chemicals and their respective functional roles (e.g., Emollient, Humectant, Preservative) along with their weights in percentages or ranges.

        ### Output Format:
        Provide the response in a simple text format with headings and bullet points, as follows:

        Functional Role: <Role Name>
        - Chemical: <Chemical Name>, Weight: <Percentage or Range>
        - Chemical: <Chemical Name>, Weight: <Percentage or Range>

        Functional Role: <Role Name>
        - Chemical: <Chemical Name>, Weight: <Percentage or Range>

        If no relevant information is found, respond with a message saying **"No New Functional roles found"** exactly and nothing else. Avoid repeating information already extracted in "Functional Role Extracted Till Now." Do not provide any comments, notes, or suggestions. Only provide the extracted data in the specified format.
        """

        # Format prompt and injecting contextual data to improve answer quality to include previous context and the current chunk
        output_format = """
        {
        "patent": {
            "patent_no": "OA06243",
            "inventor_name": ["name1", "name2", ....],
            "assignee_information": "Loreal",
            "cpcc_codes": ["Code1", "code2", "code3", "code4"]
        }
        }
        """

        document_prompt = f"""
        ### New Chunk:
        {chunk}

        ### Instruction:
        You are a data scientist working for a company that is building a report for a cosmetic patent document. Summarize the 'New Chunk' in and extract the patent no, inventor name , assignee information and cpcc codes if available in the given chunk. If this information is not avilable then send a null entry against that field, you will provide the reponse in json format :
        {output_format}
        """

        # Log token usage for name and description
        tokens_in_prompt = num_tokens_from_string(prompt)
        tokens_information_extracted = num_tokens_from_string(information_extracted)
        print(f"Tokens sent for Chunk {i}: {tokens_in_prompt}")
        print(f"information_extracted tokens sent for Chunk {i} {tokens_information_extracted}")

        # Log token usage for functional role and description
        tokens_in_prompt_2 = num_tokens_from_string(composition_prompt)
        tokens_information_extracted_2 = num_tokens_from_string(composition_information_extracted)
        print(f"Tokens sent for Chunk {i}: {tokens_in_prompt_2}")
        print(f"information_extracted tokens sent for Chunk {i} {tokens_information_extracted_2}")

        # Process the starting chunks of patents for document details - No. of Api calls 2 maximum.
        if i < 2 : # The logic needs to be more strict so that document chunk is queries for.
            document_information_extracted = await process(document_prompt, provider)
            print(f"Document Chunk {i} processed response: {document_information_extracted}")
        
        # Process the chunk for name and description information Here - No. of Api calls ~ No. of Chunks
        processedChunk = await process(prompt, provider)
        print(f"Chunk {i} processed response: {processedChunk}")

        # Process the chunk for functional_role and ingredient analysis - No. of Api calls ~ No. of chunks
        processedChunk_2 = await process(composition_prompt, provider)
        print(f"Chunk {i} processed response: {processedChunk_2}")


        # Check if new information exists and append unique parts
        if "No new information found" not in processedChunk:
            information_extracted = information_checker(information_extracted, processedChunk)
        else:
            print(f"No new information found in Chunk {i}")

        
        # Check if the response contains "No New Functional roles found"
        if "No New Functional roles found" in processedChunk_2:
            print(f"No new functional roles found in Chunk {i}. Skipping appending to information_extracted_2.")
        else:
            # Directly append the new functional roles to the existing extracted information
            print(f"Appending new functional roles from Chunk {i}.")
            composition_information_extracted += f" {processedChunk_2.strip()}"

        # Break the loop if the token count exceeds the maximum limit
        if total_token_count > max_tokens_per_chunk:
            print("Token limit exceeded, stopping chunk processing.")
            break

    print("\nFinal Information Extracted from all the text:")
    print(information_extracted)

    print("\nFinal Composition Information Extracted from all the text:")
    print(composition_information_extracted)

    # Call extract_name_description to finalize the output
    name_description_info = await extract_name_description(information_extracted, provider)
    print("\nFinalized Product Name and Description:")
    print(name_description_info)

    # Call extract_name_description to finalize the output
    functional_roles_info = await final_composition_information(composition_information_extracted, provider)


    # Call the cleaner and get document information.
    document_info = extract_document_details(document_information_extracted)

    # Finalize extracted information
    finalized_information = {
        "patent_no": document_info.get("patent_no"),
        "inventor_names": document_info.get("inventor_name"),
        "cpcc_codes": document_info.get("cpcc_codes"),
        "assignee_information": document_info.get("assignee_information"),
        "product_name": name_description_info.get("name") if name_description_info else None,
        "description": name_description_info.get("description") if name_description_info else None,
        "functional_roles": functional_roles_info.get("functional_roles") if functional_roles_info else None
    }

    print(finalized_information)
    return finalized_information
    

### Input : 
# schema , Chunks data list which is a list of the chunks of the patent text 
# It will take the prompt with the schema encoded so we will provide the user with prompt templates , need for this - KAG, graphrag all tools are good if we have a schema in mind. ( Application part )

# output :
# A graph python pandas dataframe will be created.

# This graph dataframe will have a cost associated with it right now using open ai 

# We need the open api specifications for the above endpoint - 


# ---------------------------------------------------------------------------
#  PIPELINE WRAPPER (returns HTML path + emits progress)  
# ---------------------------------------------------------------------------
ProgressCB = Callable[[str, int], None]  # (stage, percent) → None


# Then we will call the graph creation workflow , and which will take the redundant graph and then call the De-duplication process

# Once the graph has been de-duplicated according to the workflow then it will display the monte-carlo report if possible not too much work

# And then it will choose the optimal csv file create a graph using that selected file , for the purpose of this demo 

# we will choose the graph with maximum comprehension levels , we could choose one it like 20 different versions of the graph 
# without any data loss , the reader must appreciate this idea and is novel in its statement as of 2025 , no other framework does this 
# in such a manner.

# Then the final part which chat with this graph , we will simply chat with the graph using cypher query frameworks , right now we have

# preloaded cypher templates , and an LLM classifies to choose the best one , but as the system grows one can see that this can easily turn 

# into a tree of knowledge full of different cyher queries which satisfy different of application of this ground truth data.

# ─── utils/refined_graph.py ──────────────────────────────────────────────
import json, re, pandas as pd, networkx as nx
from pathlib import Path
from typing import Union

def refined_df_to_graph_html(
    df: pd.DataFrame,
    patent_number: str,
    html_name: str = "Refined-Graph.html",
) -> Path:
    """
    Build a graph in which **one central `patent_number` node** has
    `CONTAINS` edges to every representative_member.  Each representative
    keeps its `has_entry` edges to split entries.
    """
    rows = []

    # central patent node <= no properties except type/name
    patent_node_json = json.dumps({"name": patent_number, "type": "patent"}, ensure_ascii=False)

    for _, row in df.iterrows():
        rep = str(row.get("representative_member", "")).strip()
        if not rep:
            continue

        # edge: patent --contains--> representative
        rows.append(
            {
                "node_1": patent_node_json,
                "node_2": json.dumps({"name": rep, "type": "representative"}, ensure_ascii=False),
                "edge": "contains",
            }
        )

        # handle split entries
        raw = row.get("split_entries", "")
        entries = (
            [x.strip() for x in re.split(r",|;|/|\\n", raw) if x.strip()]
            if isinstance(raw, str)
            else [str(x).strip() for x in raw if str(x).strip()]
        )

        for ent in entries:
            rows.append(
                {
                    "node_1": json.dumps({"name": rep, "type": "representative"}, ensure_ascii=False),
                    "node_2": json.dumps({"name": ent, "type": "entry"}, ensure_ascii=False),
                    "edge": "has_entry",
                }
            )

    dfg = pd.DataFrame(rows, columns=["node_1", "node_2", "edge"])
    G = build_knowledge_graph(dfg)

    dst = Path(html_name).resolve()
    create_pyvis_graph(G, dst.as_posix())
    return dst

def parse_entries(x):
    if isinstance(x, list):
        return x
    if not isinstance(x, str):
        return []
    x = x.strip()
    # 1) if it looks like JSON, try json.loads
    if x.startswith("[") and x.endswith("]"):
        try:
            return json.loads(x)
        except json.JSONDecodeError:
            pass
    # 2) otherwise, split on commas
    #    or just wrap single string in a list
    return [item.strip() for item in x.split(",") if item.strip()]

# ─── right after you define neo4j_connection etc. ────
EMBED_LOCK = Lock()                # cheap concurrency guard
Embedder   = SimpleEmbedder()      # starts with **no records**

def autograph_ai_pipeline(
    chunks: List[str],
    prompt: Optional[str] = None,
    provider: str = "openai",
    progress_cb: Optional[ProgressCB] = None,
) -> str:
    """
    Heavy lifting – runs the whole pipeline and streams progress via progress_cb.
    Returns the absolute on‑disk path of the PyVis HTML graph (for CLI use).
    """

    # --------------------------------------------------
    # 1) helper that forwards any extra kwargs
    # --------------------------------------------------
    def emit(stage: str, pct: int, **extra):
        if progress_cb:
            progress_cb(stage, pct, **extra)

    # --------------------------------------------------
    # 2) build KG + save PyVis HTML
    # --------------------------------------------------
    emit("Building DataFrame", 3)
    df = documents2Dataframe(chunks)

    emit("Running LLM on chunks", 6)
    parallel_graph_list = df2Graph(df, concurrency_level=10)

    emit("Assembling edges", 12)
    dfg1 = graph2Df(parallel_graph_list)
    dfg1.replace("", np.nan, inplace=True)
    dfg1.dropna(subset=["node_1", "node_2", "edge"], inplace=True)
    dfg1["count"] = 4

    emit("Building NetworkX graph", 20)
    G = build_knowledge_graph(dfg1)

    # -------- save HTML inside /api/src --------
    from pathlib import Path
    BASE_DIR       = Path(__file__).resolve().parent   # -> /api/src
    html_filename  = "Seed-Graph.html"
    html_abs_path  = BASE_DIR / html_filename          # /api/src/Seed-Graph.html
    html_url_path  = f"/static/src/{html_filename}"    # URL the frontend can fetch

    emit("Exporting PyVis HTML", 30)
    create_pyvis_graph(G, str(html_abs_path))

    # one SSE message that includes the URL
    emit(
        "HTML ready",
        30,
        file_url=html_url_path
    )

    # --------------------------------------------------
    # 3) rest of your pipeline – unchanged logic
    # --------------------------------------------------
    emit("Finding Unique Ingredients", 35)
    unique_ing_dict = get_active_ingredients_transformed(dfg1)

    emit("Computing chunk embeddings", 50)
    CHUNK_DATA = precompute_embeddings(chunks, concurrency=5)

    emit("running dummy test on chunk embeddings", 55)
    relevant_chunks = find_relevant_chunks("methyl paraben", CHUNK_DATA, 7)
    print(f"{len(relevant_chunks)} chunks found")
    print(relevant_chunks)

    emit("First Refinement", 65)
    refined_df, docx, csv = process_all_ingredients_simple(
        unique_ing_dict,
        CHUNK_DATA,
        output_docx="refined_ingredients_1st_cycle.docx",
        output_csv="refined_ingredients_1st_cycle.csv",
        take_user_input=False
    )

    emit("Monte‑Carlo deduplication and threshold selection", 75)
    full_df, _, _ = monte_carlo_analysis(refined_df, "MC_outputs")

    monteCarloPlotName = "MC_outputs/combined_plot.png"
    monteCarloPlot_url_path = f"/static/src/{monteCarloPlotName}"
    emit(
        "Monte Carlo Plot Ready",
        80,
        file_url=monteCarloPlot_url_path
    )

    if full_df is None:
        raise RuntimeError("No Full‑Analysis threshold satisfied ‘no repeats’ – check the summary CSV.")

    emit("2nd Refinement Cycle to build the refined KG", 80)
    final_df = refine_deduplicated_dataframe(full_df, CHUNK_DATA, False)

    
    csv_name     = "refined_df_latest.csv"
    csv_abs_path  = BASE_DIR / csv_name 
    csv_url = f"/static/src/{csv_name}"
    final_df.to_csv(csv_abs_path, index=False)

    emit("Refined CSV ready",               # let the frontend know
     83,
     file_url=csv_url)

    # --------------------------------------------------
    #  (RE)BUILD the semantic‑search index from that CSV
    # --------------------------------------------------
    with EMBED_LOCK:
        df_embed = pd.read_csv(csv_abs_path)
        # you can customise what constitutes “text”
        rows = []
        for _, r in df_embed.iterrows():
            entries = parse_entries(r["split_entries"])
            text = f"{r['representative_member']} " + " ".join(entries)
            rows.append({
                "text":   text,
                "labels": ["refined"],
                "name":   r["representative_member"][:60],
            })

        Embedder.records.clear()
        Embedder.emb = None
        Embedder.add_records(rows)
        print(f"[semantic] re-indexed {len(rows)} refined rows")


    # ▲▲▲  NEW – build & expose the refined graph  ▲▲▲
    html_refined = refined_df_to_graph_html(final_df, "DummyPatent", "Refined-Graph.html")
    emit("Refined KG ready",                 # SSE for the frontend
         82,
         file_url=f"/static/src/{html_refined.name}")
    # ▲▲▲  end of insertion  ▲▲▲

    emit("Ground Truth Application – Detecting similar products from yesstyle.com", 90)
    reports = run_analysis_for_multiple_csvs(
        final_df,
        product_csv_path="products.csv",
        default_only_claim=DEFAULT_ONLY_CLAIM,
        default_comprehension_level=DEFAULT_COMPREHENSION_LEVEL,
        start_row_num=1,
        end_row_num=6340
    )

    for matched_path, summary_path in reports:
        # 2) matched‑products CSV
        emit("Report ready", 91,
            file_url=f"/static/{Path(matched_path).name}")

        # 3) summary CSV
        emit("Report ready", 92,
            file_url=f"/static/{Path(summary_path).name}")

    topkfilename = "products_topk_matched_number_1.csv"
    topkfile_url_path = f"/static/src/{topkfilename}" 
    emit(
        "Topk File ready",
        94,
        file_url=topkfile_url_path
    )
    # ----------------------------------------------------------
    # NEW ❱  build + expose *Top‑10* as tiny JSON for frontend
    # ----------------------------------------------------------
    try:
        # ➊ read the just‑written top‑K CSV
        df_topk        = pd.read_csv(topkfilename).head(10)

        # ➋ keep only the 3 fields the React modal needs
        top10_json_rec = df_topk[["product_id",
                                  "product_url",
                                  "match_percentage"]]

        json_name      = "products_top10_number_1.json"
        top10_json_path = Path(json_name)            # on disk
        top10_json_rec.to_json(top10_json_path,
                               orient="records",
                               force_ascii=False)

        # ➌ push the URL to the client via the usual SSE ‘emit’
        emit("Top‑10 JSON ready",
             96,
             file_url=f"/static/src/{json_name}")

    except Exception as e:     # never abort the pipeline
        print(f"[WARN] Top‑10‑JSON step failed → {e}")

    emit("Generated final reports", 97)


    emit("Done", 100)
    return str(html_abs_path)

######################################################################################
# Description : Autograph AI Product-Patent Matching Pipeline
# Author : Kanishk Mittal
# License : Apache 2.0
######################################################################################
import sys
from yachalk import chalk
sys.path.append("..")

import openai
import json

# Helpers

import uuid

import tiktoken  # <-- NEW: we use this library for more accurate token counting
import numpy as np
import psutil
import multiprocessing
import platform
from concurrent.futures import ThreadPoolExecutor, as_completed

## Knowlegde Graph creation
import re
import networkx as nx
import json
import pandas as pd
from pyvis.network import Network
from networkx.algorithms import community
import seaborn as sns
import random
import os

openai.api_key = openai_key

def graphPrompt(input: str, metadata={}, model="gpt-4o"):
    if model == None:
        model = "gpt-4o"

    CLAIM_PROMPT = (
             """
             You are a network-graph maker who extracts terms, their relationships, and relevant attributes from a given context.

            The context is delimited by triple backticks. Your tasks are:
            
            "Thought 1": While traversing through each sentence or relevant text segment within the patent text, identify key terms. 
            - Terms are any concepts or entities essential to the central meaning of the context. 
            - You must pay special attention to chemical substances (e.g., “active_ingredient”, “chemical_compound”, “functional_group”, “pharmaceutical_salt”, “excipient”, “polymer”), but you can also extract other relevant terms that appear crucial (e.g., “apparatus/device”, “process”, “organisation”, “condition”, “place”, “event”, “concept”, “misc”). 
            - Keep these terms **atomistic**: avoid merging distinct sub-parts into a single node. For example, a polymer name should be kept separate from any closely related derivative name if the patent text treats them distinctly.
            
            "Thought 2": Determine possible pairwise relationships between the extracted terms you have to cover all the chemicals dont miss any. For instance:
            - **hasFormula**: for a situation where a “compound #” is disclosed along with its structural or chemical name (e.g., “Compound 1” → “(Chemical name)”), or a polymer with a more specific structure described.  
            - **derivativeOf**, **coOccursWith**, or any similarly concise relationship (≤ 3 words) that the text implies (but be consistent; do not invent synonyms).
            
            "Thought 3": For each extracted term (node), you must store **two main booleans** (which can both be `true` if appropriate):
            - `is_active_ingredient`
            - `is_structural_component`
            
            Alongside each boolean, you should record a **weight** field if applicable:
            - `active_ingredient_weight`: Float or range (if `is_active_ingredient = true`).
            - `structural_component_weight`: Float or range (if `is_structural_component = true`).
            - `standard_names`: If the associated text provide example or standard names in case the ingredient name is a brand or marketed name and not the original INCI name ,then you will provide the synonyms of the ingredient name present in the text , and also from your own knowledge about cosmetic chemicals
            - `functional_role`: Select only one from this list functional_roles dont choose your own,  = 
                          ["abrasive", "absorbent", "anti-acne agent", "anti-aging agent", 
                            "anti-caking agent", "anti-dandruff agent", "anti-pollution agent", Anti-inflammatory agent
                            "anticorrosive agent", "antifoaming agent", "antimicrobial agent", "antioxidant", 
                            "antiperspirant", "antiplaque agent", "antistatic agent", "astringent", 
                            "binding agent", "bleaching agent", "bulking agent", "chelating agent", 
                            "cleansing agent", "collagen booster", "colorant", "decontracting (dermo-relaxing) agent", 
                            "denaturant", "deodorant", "depigmenting agent", "depilatory agent", "emollient", 
                            "emulsifier", "emulsion stabilizer", "film forming agent", "fibroblast proliferation agent", 
                            "flavoring agent", "foam booster", "fragrance", "hair conditioning agent", 
                            "hair dyeing agent", "hair fixative", "humectant", "hydrotrope", "keratolytic agent", 
                            "keratinocyte proliferation agent", "light stabilizer", "masking agent", 
                            "nail conditioning agent", "opacifier", "oral care agent", "oxidizing agent", 
                            "pearlescent agent", "plasticizer", "preservative", "propellant", "reducing agent", 
                            "skin brightening agent", "skin conditioning agent", "skin protectant", "slimming agent", 
                            "solubilizer", "solvent", "soothing agent", "surfactant", "sweetener", "tonic", 
                            "UV absorber", "UV filter", "viscosity modifier", "wetting agent", "pH adjuster", 
                            "anti-wrinkle agent"]
            
            If the context explicitly states that a substance is an active ingredient in a composition (e.g., a brightening agent), then set `is_active_ingredient = true`. If the context indicates 
            it is just part of a carrier, backbone, or generally a structural or excipient component, then `is_structural_component = true`. In many patent claims, both might apply (e.g., the same 
            substance can serve multiple roles) — or neither if it’s not used in a direct compositional function.
            
            Whenever a weight or range (e.g., “0.5–3% by weight”) is mentioned for that substance, record it in the corresponding `_weight` field.  
            If no relevant function or weight is given, keep both booleans `false` and omit weight fields.  
            
            In all cases, you must add an `"explanation"` field of **at least 100 words** describing the rationale for your assignment of booleans (why or why not an active ingredient or structural 
            component), referencing at least **2–3 keywords or short phrases** from the immediate context so that a human can verify the correctness.
            
            "Thought 4": Create a **WO2023060387** (or an analogous single node that represents the entire composition or claim). For any node identified as an active or structural ingredient, 
            link it to this composition_node using the relationship `"contains"`.  

            - The graph should have all the components connected by this belongs_to edge , all the ingredients which I am giving you belong to the composition
            
            - If a node is neither an active ingredient nor a structural component, do **not** link it to the composition_node.
            
            "Thought 5": **Output** must be a JSON **list** of objects, each capturing a single triplet in the graph. The structure is:
                    [
                      {
                        "node_1": {
                          "name": "First concept extracted",
                          "is_active_ingredient": true or false,
                          "active_ingredient_weight": Float or range
                          "is_structural_component": true or false 
                          "structural_component_weight": Float or range
                          "explanation": "Focus on the Why and give proper  What, Why with The contextual complete is this node a relevant ingredient and if structural_component or an active_ingredient or both or none then why ? ( Minimum 100 words) Crisp solid justified explanation of why do you think its True or False. You must also include the relevant text snippet or 2–3 keywords from the chunk to anchor your reasoning, and so that humans can verify your reasoning.\n".",
                          "standard_name": "Choose the standard INCI name from your knowlegde or the text provided , otherwise mention - 'not found' if the standard name is not found"
                          "functional_role: "Choose one from the list of INCI functional roles made available to you"
                          "examples_mentioned": suppose the text says that the composition contains cross linked copolymer , it will provide example of which polymers or chemicals to use for that category fetch all those example and create a list"
                        },
                        "node_2": {
                          "name": "Second concept extracted",
                          "is_active_ingredient": true or false,
                          "active_ingredient_weight": Float or range
                          "is_structural_component": true or false 
                          "structural_component_weight": Float or range
                          "explanation": "Focus on the Why and give proper what, why with The contextual complete is this node a relevant ingredient , crisp solid justified explanation of why do you think its True or False. ( Minimum 100 words) Crisp solid justified explanation of why do you think its True or False. You must also include the relevant text snippet or 2–3 keywords from the chunk to anchor your reasoning, and so that humans can verify your reasoning.\n".",
                          "standard_name": "Choose the standard INCI name from your knowlegde or the text provided , otherwise mention - 'not found' if the standard name is not found"
                          "functional_role: "Choose one from the list of INCI functional roles made available to you"
                          "examples_mentioned": suppose the text says that the composition contains cross linked copolymer , it will provide example of which polymers or chemicals to use for that category fetch all those example and create a list"
                        },
                        
                        "edge": "Short relationship (<=3 words)"
                      },
                      ...
                    ]

           """
    )

    SYS_PROMPT = (
     
        """
         You are a network graph maker who extracts terms, their relations, and relevant attributes from a given context.
        
                The context is delimited by triple backticks. Your tasks are:
                
                "Thought 1: While traversing through each sentence, Think about the key terms mentioned in it.\n"
                    "Terms are basically 'concepts' or 'entities' that are essential to the central meaning of context, with special emphasis on these- 
                    "active_ingredient\", \"chemical_compound\", \"functional_group\", \"pharmaceutical_salt\", \"excipient\",
                    \"polymer\", \"process\", \"apparatus / device\", \"document\", \"organisation\", \"condition\", \"place\" ,
                    \"event\",\"concept\" ,\"misc\"] \n"
                    
                    "\tTerms should be as atomistic as possible and avoid merging distinct sub-parts into a single node."
        
                "Thought 2": Think about how these terms can have one on one relation with other terms. 
                    **Some Important relationships you can expect ** between the nodes, focusing on:
                           - 'hasFormula' (especially for “compound #” → “(chemical name or moiety)”), 
                           - 'coOccursWith',
                           - 'saltFormOf',
                           - 'derivativeOf',
                           - 'fusedRingWith',
                           - or any succinct relation that the context implies (≤ 3 words), but be consistent don't introduce redudant relations that are synonyms of each other. Be consistent
                            Example of consistency - 
                            If you see something like “(3,5-dibromo-4-hydroxy-phenyl)-(2,3-dihydro...) (compound 1),” create two nodes:
                             • node_1 = “compound 1”
                             • node_2 = “(3,5-dibromo-4-hydroxy-phenyl)-(2,3-dihydro...) (compound 1)”
                             and set `"edge": "hasFormula"` or a similarly relevant relation.
        
                             Then for all such nodes compound 2 , 3 four you will use the same relation "hasFormula" to show their formula, and stay consistent on this relationship.
                           
                    
                "Thought 3": Here , You must think about this based on the node and relationship values and then choose. 
                 For each node, you will store TWO booleans, and they both CAN be true simultaneously:
                   - 'is_active_ingredient': true/false
                   - 'active_ingredient_weight': Float or range
                   - 'is_structural_component': true/false
                   - 'structural_component_weight': Float or range
                   
                   If the node is the main brightening/whitening agent, **Check if a node is a relevant ingredient** or functionally 
                   significant and then set is_active_ingredient=true and is_structural_component=false and then check its weight specified 
                   then accurately copy it in the active_ingredient weight field if is_active_ingredient=true and structural_component_weight field if is_structural_component_=true 
                   and provide explanation 
                   
                   If you have strong evidence from the context that a compound is a relevant ingredient, set 
                   'is_active_ingredient' = true and is_structural_component=false and then check its weight specified 
                   then accurately copy the weight found in the active_ingredient_weight field if is_active_ingredient=true and provide explanation 
                   
                   If you have strong evidence from the context that the node is just a sub-part of a formula (like a ring, moiety, or 
                   partial chemical structure), then is_structural_component=true and is_active_ingredient=false and and then check its weight specified 
                   then accurately copy the weight found in the active_ingredient_weight field if is_structural_component_=true provide explanation 
                   If neither condition applies, keep both false and no weights need to be added provide explanation   
                   if Both Condition applies, keep both True and provide weight values in both the weight fields accurately provide one explanation for both
                   For each node, if you believe it is an 'active ingredient' or an 'is_structural_component' or both or none, then 
                   you will set the respective 'is_active_ingredient' or 'is_structural_component' fields and weight fields
                   and then you MUST MANDATORILY provide a justification in 'explanation' of your choices at least 100 words justifying 
                   your choice.You must also include the relevant text snippet or 2–3 keywords from the chunk to anchor your reasoning, and 
                   so that humans can verify your reasoning.
        
                "Thought 4": Finally, after deciding which nodes qualify as active ingredients or structural components, you should create a special relationship 
                called "is_part_of_composition_by_weight" between each qualifying node and a central composition_node.
        
                    Why do we have a composition_node?
                    
                    This composition_node acts as a super node in the graph. It’s used as the main entry point, connecting to every node that is recognized as 
                    either an active ingredient or a structural component.
                    How do we connect nodes?
                    
                    For any node identified as an active ingredient or a structural component (or both), link it to the composition_node with the 
                    "is_part_of_composition_by_weight" relationship.
                    What about nodes that don’t qualify?
                    
                    If a node is neither an active ingredient nor a structural component, it should not be connected to the composition_node.
                    The result is a graph where the composition_node is highly central, because it’s directly connected to all nodes representing active 
                    ingredients and/or structural components. This approach makes it easy to trace how each element in your system contributes to the 
                    overall composition.
                                
        
                
                "Thought 4". **Output Format**: Return a **list** of JSON objects, each containing a pair of nodes and their attributes: "node_1": "First concept extracted from the text", "node_2": "Second concept extracted from the text",
                
                    [
                      {
                        "node_1": {
                          "name": "First concept extracted",
                          "is_active_ingredient": true or false,
                          "active_ingredient_weight": Float or range
                          "is_structural_component": true or false 
                          "structural_component_weight": Float or range
                          "explanation": "Focus on the Why and give proper  What, Why with The contextual complete is this node a relevant ingredient and if structural_component or an active_ingredient or both or none then why ? ( Minimum 100 words) Crisp solid justified explanation of why do you think its True or False. You must also include the relevant text snippet or 2–3 keywords from the chunk to anchor your reasoning, and so that humans can verify your reasoning.\n",
                          "standard_name": "Choose the standard INCI name from your knowlegde or the text provided , otherwise mention - 'not found' if the standard name is not found",
                          "functional_role: "Choose one from the list of INCI functional roles made available to you",
                          "category": "Choose from the categories above"
                        },
                        "node_2": {
                          "name": "Second concept extracted",
                          "is_active_ingredient": true or false,
                          "active_ingredient_weight": Float or range
                          "is_structural_component": true or false 
                          "structural_component_weight": Float or range
                          "explanation": "Focus on the Why and give proper what, why with The contextual complete is this node a relevant ingredient , crisp solid justified explanation of why do you think its True or False. ( Minimum 100 words) Crisp solid justified explanation of why do you think its True or False. You must also include the relevant text snippet or 2–3 keywords from the chunk to anchor your reasoning, and so that humans can verify your reasoning.\n",
                          "standard_name": "Choose the standard INCI name from your knowlegde or the text provided , otherwise mention - 'not found' if the standard name is not found",
                          "functional_role: "Choose one from the list of INCI functional roles made available to you",
                          "category": "Choose from the categories above"
                        },
                        
                        "edge": "Short relationship (<=3 words)"
                      },
                      ...
                    ]
                Some Rules to Keep in Mind - 
                5. **Do not add extra keys** beyond these:
                   - node_1
                   - node_2
                   - edge
                
                6. **Use only text evidence** from the context to decide attributes; do not invent or guess beyond what is stated.
        
                7. **Preserve uniqueness**: If two expressions are exact synonyms referring to the very same entity (not just similar), 
                unify them in the same node rather than repeating them. If they are partial synonyms but the text implies they are different 
                (e.g., distinct moieties or derivatives), keep them separate.
        
                Now, analyze the text below and produce the JSON list of relations strictly in the format described:
                """
    )

    USER_PROMPT = f"context: ```{input}``` \n\n output: "
    # Call the OpenAI ChatCompletion endpoint with system and user messages:
    response = openai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": CLAIM_PROMPT},
            {"role": "user", "content": USER_PROMPT},
        ],
        temperature=0.7,
    )

    # Get assistant's text output:
    assistant_message = response.choices[0].message.content
    # Optional: remove backticks
    assistant_message = assistant_message.replace("```json","").replace("```","")
    try:
        result = json.loads(assistant_message)
        print("\n\nSUCCESS ### Here is the correct response: ", assistant_message, "\n\n")
        result = [dict(item, **metadata) for item in result]
    except:
        print("\n\nERROR ### Here is the buggy response: ", assistant_message, "\n\n")
        result = None
    return result



### Helpers
def documents2Dataframe(documents) -> pd.DataFrame:
    rows = []
    for chunk in documents:
        if isinstance(chunk, str):
            row = {
                "text": chunk,
                "chunk_id": uuid.uuid4().hex,
            }
        else:  # assume LangChain Document
            row = {
                "text": chunk.page_content,
                **chunk.metadata,
                "chunk_id": uuid.uuid4().hex,
            }
        rows.append(row)
    df = pd.DataFrame(rows)
    return df

def df2Graph(dataframe: pd.DataFrame, concurrency_level=4) -> list:
    total = len(dataframe)
    results = []

    def process_chunk(i, row):
        print(f"Processing chunk {i+1}/{total} (chunk_id={row['chunk_id']})...")
        result = graphPrompt(row["text"], {"chunk_id": row["chunk_id"]}) 
        # for i, item in enumerate(result):
        #     for k, v in item.items():
        #         print(f"The key is : {k} and value is : {v}")
        return (row["chunk_id"], result)

    futures = {}
    with ThreadPoolExecutor(max_workers=concurrency_level) as executor:
        for i, row in dataframe.iterrows():
            future = executor.submit(process_chunk, i, row)
            futures[future] = row["chunk_id"]

        for future in as_completed(futures):
            chunk_id = futures[future]
            try:
                chunk_id_result, data = future.result()
                results.append(data)
                # print(f"Futured Chunk Data :  {data} , {type(data)}")
                # print("Chunk futured successfully")
            except Exception as exc:
                print(f"Chunk {chunk_id} generated an exception: {exc}")

    # Build valid_results using a simple for loop
    valid_results = []
    for r in results:
        if r is not None:
            valid_results.append(r)

    # Flatten valid_results into concept_list using nested for loops
    concept_list = []
    for sublist in valid_results:
        for item in sublist:
            concept_list.append(item)

    return concept_list


def graph2Df(nodes_list) -> pd.DataFrame:
    """
    Build a DataFrame from a list of dicts, each having keys 'node_1', 'node_2', 'edge'.
    'node_1' and 'node_2' are dicts of node properties, so we convert them to JSON strings
    to store in CSV. 'edge' is a short string. Return DF with exactly 3 columns.
    """
    graph_dataframe = pd.DataFrame(nodes_list)

    # Ensure we have "node_1", "node_2", and "edge"
    def to_json_str(obj):
        if isinstance(obj, str):
            # Try to parse it first—if that fails, it means it wasn't valid JSON anyway.
            try:
                # parse the string into a dict
                parsed = json.loads(obj)
                # now do a proper dump
                return json.dumps(parsed, ensure_ascii=False)
            except:
                # if it wasn’t valid JSON, just return the string
                return obj
        elif isinstance(obj, dict):
            return json.dumps(obj, ensure_ascii=False)
        else:
            return str(obj) if obj is not None else ""

    if "node_1" in graph_dataframe.columns:
        graph_dataframe["node_1"] = graph_dataframe["node_1"].apply(to_json_str)
    if "node_2" in graph_dataframe.columns:
        graph_dataframe["node_2"] = graph_dataframe["node_2"].apply(to_json_str)

    if "edge" not in graph_dataframe.columns:
        graph_dataframe["edge"] = ""

    # Keep only these three columns
    graph_dataframe = graph_dataframe[["node_1", "node_2", "edge"]]

    return graph_dataframe



def build_knowledge_graph(dfg1):
    G = nx.Graph()
    for idx, row in dfg1.iterrows():
        node1_props = json.loads(row["node_1"])
        node2_props = json.loads(row["node_2"])
        node1_name = node1_props["name"]
        node2_name = node2_props["name"]
        G.add_node(node1_name, **node1_props)
        G.add_node(node2_name, **node2_props)
        edge_label = row["edge"]
        G.add_edge(node1_name, node2_name, label=edge_label, count=row.get("count", None))
    return G

def colorize_communities(G, net):
    communities_gen = community.girvan_newman(G)
    top_level = next(communities_gen)
    comm_list = sorted(map(sorted, top_level))

    palette = sns.color_palette("hls", len(comm_list)).as_hex()
    random.shuffle(palette)
    
    for i, comm_nodes in enumerate(comm_list):
        color_hex = palette[i]
        for node_id in comm_nodes:
            for n in net.nodes:
                if n["id"] == node_id:
                    n["color"]["background"] = color_hex
                    break

def create_pyvis_graph(G, output_html="graph.html"):
    net = Network(
        notebook=False,
        height="700px",
        width="100%",
        cdn_resources="remote",
        select_menu=True,
        filter_menu=False
    )
    
    net.from_nx(G)

    # Transform labels to include newlines & configure multi-line support
    for node in net.nodes:
        # Replace spaces with line breaks (or break after ~12 chars, etc.)
        node["label"] = re.sub(r"\s+", "\n", node["label"])
        
        node["size"] = 60
        node["shape"] = "circle"
        node["borderWidth"] = 2
        
        # IMPORTANT: multi=True allows \n to be interpreted as line breaks
        node["font"] = {
            "size": 30,
            "face": "arial",
            "color": "#000000",
            "multi": True
        }
        node["color"] = {
            "background": "#A9CCE3",
            "border": "#2471A3",
            "highlight": {
                "background": "#F7DC6F",
                "border": "#B7950B"
            }
        }

    for edge in net.edges:
        edge["width"] = 2
        edge["color"] = "#848484"
        # edge["smooth"] = {"type": "continuous"} # Uncomment for curved edges

    colorize_communities(G, net)

    # Use Barnes-Hut with 'avoidOverlap' to keep nodes spaced
    net.set_options("""
    var options = {
      "physics": {
        "barnesHut": {
          "gravitationalConstant": -50000,
          "centralGravity": 0.2,
          "springLength": 200,
          "springConstant": 0.04,
          "avoidOverlap": 1
        },
        "minVelocity": 0.75
      }
    }
    """)

    net.show(output_html, notebook=False)

def get_active_ingredients_transformed(dfg1):
    """
    Collects nodes where 'is_active_ingredient' or 'is_structural_component' is True.
    Each unique ingredient name (lower-cased) becomes a key in 'final_dict', with its value a dictionary:
      {
        "standard_names": [ ... ],
        "functional_roles": [ ... ],
        "is_active_ingredient": Boolean,
        "is_structural_component": Boolean,
        "examples_mentioned": [ ... ]
      }
    Returns 'final_dict' and also prints 'active_ingredient_list' (a list of key-value pairs).
    """
    final_dict = {}

    for idx, row in dfg1.iterrows():
        for col in ["node_1", "node_2"]:
            try:
                node_data = json.loads(row[col])  # parse JSON from the cell

                # Only track ingredients that are active or structural
                if node_data.get("is_active_ingredient") or node_data.get("is_structural_component"):
                    ing_name = node_data["name"].strip().lower()

                    # If this ingredient hasn't been seen, initialize its record
                    if ing_name not in final_dict:
                        final_dict[ing_name] = {
                            "standard_names": set(),       # Use sets to avoid duplicates
                            "functional_roles": set(),
                            "is_active_ingredient": False,
                            "is_structural_component": False,
                            "examples_mentioned": set()
                        }

                    # Safely add standard_name (assuming it is a string)
                    if isinstance(node_data["standard_name"], str):
                        final_dict[ing_name]["standard_names"].add(node_data["standard_name"].strip())

                    # Safely add functional_role (assuming it is a string)
                    if isinstance(node_data["functional_role"], str):
                        final_dict[ing_name]["functional_roles"].add(node_data["functional_role"].strip())

                    # Handle examples_mentioned which can be a list of strings
                    # or occasionally a single string
                    if isinstance(node_data["examples_mentioned"], list):
                        for example in node_data["examples_mentioned"]:
                            final_dict[ing_name]["examples_mentioned"].add(example.strip())
                    elif isinstance(node_data["examples_mentioned"], str):
                        final_dict[ing_name]["examples_mentioned"].add(node_data["examples_mentioned"].strip())

                    # If any occurrence sets active_ingredient or structural_component to True, keep it True
                    if node_data.get("is_active_ingredient"):
                        final_dict[ing_name]["is_active_ingredient"] = True
                    if node_data.get("is_structural_component"):
                        final_dict[ing_name]["is_structural_component"] = True

            except (json.JSONDecodeError, TypeError, KeyError):
                # Ignore or handle errors if needed
                pass

    # Convert sets to lists for final output
    for ing in final_dict:
        final_dict[ing]["standard_names"] = list(final_dict[ing]["standard_names"])
        final_dict[ing]["functional_roles"] = list(final_dict[ing]["functional_roles"])
        final_dict[ing]["examples_mentioned"] = list(final_dict[ing]["examples_mentioned"])

    # Build a list of key-value pairs
    active_ingredient_list = []
    for key, val in final_dict.items():
        active_ingredient_list.append({key: val})

    # Print everything
    print(f"Total unique entries: {len(final_dict)}\n")
    for item in active_ingredient_list:
        print(item)

    return final_dict


import openai
from typing import List, Dict, Any
import math
from rank_bm25 import BM25Okapi


#####################################################################
## 2) Embedding-based retrieval logic
#####################################################################

EMBEDDING_MODEL = "text-embedding-ada-002"  # or your preferred model

def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Compute the cosine similarity between two vectors."""
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    mag1 = math.sqrt(sum(a * a for a in vec1))
    mag2 = math.sqrt(sum(b * b for b in vec2))
    if not mag1 or not mag2:
        return 0.0
    return dot_product / (mag1 * mag2)

def embed_text(text: str) -> List[float]:
    """Call OpenAI's embedding API and return the embedding vector."""
    try:
        print("The text being sent is :", text)
        response = openai.embeddings.create(
            model="text-embedding-3-large",
            input=text
        )
        # 'data' is a list of embeddings; we only sent one input, so it's 0th
        return response.data[0].embedding
    except Exception as e:
        print(f"[Error] Embedding failed for text: {text[:30]}... => {e}")
        return []



def is_specific_chemical(ingredient_name: str) -> bool:
    """
    Determine via LLM if the ingredient name is a specific chemical compound or a generic term.
    Returns True if specific (discrete chemical), or False if generic/broad.
    """
    # Prompt the LLM to answer strictly with True/False
    query = (f"Is '{ingredient_name}' the name of a specific chemical substance (a single defined compound)? "
             "Answer with True or False only.")
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",  # using GPT-4o for quick classification (or use gpt-4 if desired)
            messages=[{"role": "user", "content": query}],
            temperature=0.7,
        )
        answer = response.choices[0].message.content
    except Exception as e:
        print(f"[Error] LLM classification failed for {ingredient_name}: {e}")
        return False  # default to False (generic) on error
    
    # Interpret the answer robustly
    answer_lower = answer.lower()
    if "true" in answer_lower:
        return True
    elif "false" in answer_lower:
        return False
    else:
        # If the LLM gave an unexpected response, decide based on heuristic (e.g., presence of space suggests generic category)
        return False

def refine_ingredient_data_via_llm(original_data: Dict[str, Any], relevant_chunks: List[Any], take_user_input: bool = True) -> Dict[str, Any]:
    """
    1) Takes in the 'original_data' describing the ingredient (like from unique_ing_dict):
       {
         'original_ingredient_name': str,
         'original_standard_names': list,
         'original_functional_role': list,
         'originally_an_active_ingredient': bool,
         'originally_a_structural_component': bool
         ... anything else ...
       }

    2) Takes the 'relevant_chunks' (the top 3 or so strings/documents) as context

    3) Builds a prompt to refine the data and produce the JSON with fields:
       {
         "LLM_refined_name": "...",
         "LLM_refined_standard_names": [...],
         "LLM_refined_source_supplier": [...],
         "LLM_refined_supplying_company": [...],
         "LLM_refined_functional_roles": [...],
         "LLM_refined_mentioned_examples": [...],
         "LLM_refined_is_active_ingredient": true/false,
         "LLM_refined_is_structural_ingredient": true/false
       }

    4) Returns the parsed JSON or empty if parse fails
    """

    # Step A: Display the original data in a "debugger console"
    print("\n==========================")
    print("DEBUGGER: Original Ingredient Data")
    print("==========================")
    for k,v in original_data.items():
        print(f"{k}: {v}")

    # Step B: Compose the context from the top relevant chunks
    chunk_texts = []
    for idx, ch in enumerate(relevant_chunks):
        if hasattr(ch, "page_content"):
            # It's a Document
            chunk_texts.append(f"---Relevant Chunk #{idx+1}---\n{ch.page_content}")
        else:
            chunk_texts.append(f"---Relevant Chunk #{idx+1}---\n{str(ch)}")

    combined_context = "\n\n".join(chunk_texts)

    # Step C: Build the LLM prompt
    # We feed the original data + chunk text, ask for refined JSON
    system_instructions = (
        "You are an assistant that refines cosmetic ingredient data using the text context."
        "Return a strict JSON with the specified keys only. If something is not found, use empty or false."
    )

    # Original data as JSON
    original_data_json = json.dumps(original_data, ensure_ascii=False)

    user_prompt = f"""Here is the original data for an ingredient:
        ORIGINAL_DATA = {original_data_json}
        
        Below are the top relevant text chunks about this ingredient:
        {combined_context}
        
        From the original data + relevant text:
        1) Mandatorily, refine the 'original_ingredient_name' to a more context-specific name , the best fit based on the context.
        Refinement basically means selecting a more specific ingredient highlighted like mineral salt but which one is most preferred
        2) Derive or confirm standard names, suppliers, functional roles, examples, whether it's active or structural, etc.
        3) Important: If placeholders (e.g. “X,” “R,” “S”) appear in chemical formulas, treat them purely as variable/structural positions. 
        Do not interpret them as actual ingredient names, brand names, or functional substances. Only map them to the actual chemical 
        definitions/purposes if the text explicitly describes them as a standalone ingredient. Otherwise, keep them as placeholders.


        Return JSON ONLY with these fields:
        {{
          "LLM_refined_name": "the refined standard name",
          "LLM_refined_name_placeholder_for": if original name is a placeholder map them to the structural position in the formula accurately and 
          if they are actual active ingredients , brand names or functional substances then map them in this field according , can be a list 
          "LLM_refined_role_of_placeholder":If a chemical is a placeholder , you will justify in less than 10 words , how , it could a functional
          bridge between two other actives or structural component , present an overview of the placeholder detected above ( if any else leave empty)
          "LLM_refined_standard_names": fill in this can be a list [],
          "LLM_refined_source_supplier": fill in this can be a list [],
          "LLM_refined_supplying_company": fill in this can be a list [],
          "LLM_refined_functional_roles": Important to fill in this can be a list [],
          "LLM_refined_mentioned_examples": Try to fill in this can be a list  [],
          "LLM_refined_is_active_ingredient": must be there True/ false,
          "LLM_refined_is_structural_ingredient":must be there True/false
        }}
        Ensure it's valid JSON. If something is not found, keep it empty or false.
        """

    # Step D: Call the LLM
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_instructions},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3
        )
        llm_text = response.choices[0].message.content
        llm_text_1 = llm_text.replace("```json","").replace("```","")
        
        
    except Exception as e:
        print(f"[Error] refine_ingredient_data_via_llm => {e}")
        return {}

    # Step E: Parse the LLM's JSON
    refined = {}
    try:
        refined = json.loads(llm_text_1)
    except json.JSONDecodeError:
        print("[Warning] LLM did not return valid JSON. Full text:")
        print(llm_text_1)
        refined = {}

    print("\n==========================")
    print("Detected Relevant Chunks:")
    print("==========================")
    if relevant_chunks:
        print(relevant_chunks)
    else:
        print("(No relevant chunks found)")

    # Step F: Show the refined data to the user in the console
    print("\n==========================")
    print("DEBUGGER: LLM-Refined Data")
    print("==========================")
    if refined:
        for k,v in refined.items():
            print(f"{k}: {v}")
    else:
        print("(No valid refined data parsed)")

    # Step G: If user input is enabled, prompt to confirm or discard. Otherwise, auto-accept.
    if take_user_input:
        user_input = input("\nPress Enter to accept this refined data, or type 'n' to discard: ").strip().lower()
        if user_input == 'n':
            print("[Info] User discarded LLM-refined data.")
            return {}
        else:
            print("[Info] User accepted LLM-refined data.")
            return refined
    else:
        # If we are NOT taking user input, just accept automatically
        if refined:
            print("[Info] Automatically accepting LLM-refined data (take_user_input=False).")
            return refined
        else:
            print("[Info] No refined data to accept; returning empty.")
            return {}

##############################################################################
# 3) EXACT MATCH + EMBEDDING RERANK + (OPTIONAL) BM25 FALLBACK
##############################################################################

def find_relevant_chunks(
    query: str,
    chunk_data: List[Dict[str, Any]],
    top_k: int = 5,
    use_bm25_fallback: bool = True
) -> List[Any]:
    """
    1) EXACT SUBSTRING MATCH:
       - Return only chunks whose text (case-insensitive) contains `query`.
       - Then re-rank those matched chunks via embedding similarity, **using the
         already precomputed embeddings** (no extra API calls).

    2) If NO EXACT MATCHES:
       - Optionally run BM25 over the entire corpus, pick top results,
         then re-rank them by embedding similarity (again using precomputed
         chunk_data[i]["embedding"]).
       - If no embedding is found, skip or treat sim=0.

    Args:
        query (str): Substring or brand name (e.g. "COMPAGNIE DE VICHY").
        chunk_data (List[Dict[str,Any]]): Each dict has:
          {
            "id": int,  (optional)
            "text": <Document or string>,
            "embedding": <List[float]>
          }
        top_k (int): Number of final chunks to return.
        use_bm25_fallback (bool): If True, do lexical search if no substring matches.

    Returns:
        List of chunk 'text' objects in descending order of relevance.
    """

    query_lower = query.lower()

    # STEP A: EXACT SUBSTRING FILTER
    exact_matches = []
    for item in chunk_data:
        if hasattr(item["text"], "page_content"):
            raw_text = item["text"].page_content
        else:
            raw_text = str(item["text"])

        if query_lower in raw_text.lower():
            exact_matches.append(item)

    if exact_matches:
        # We found substring matches => re-rank by precomputed embeddings
        # [No additional embed calls here.]

        # Possibly embed the query itself (only once)
        # But you must do so if you want to do an embedding-based re-rank
        query_emb = None
        # If you have it precomputed for queries, skip. Otherwise, embed once:
        # For demonstration, let's do a single embed call:
        query_emb = _embed_query(query)

        scored = []
        for match_item in exact_matches:
            chunk_emb = match_item["embedding"]
            if not chunk_emb:
                # If it truly has no precomputed embedding, skip or set sim=0
                sim = 0.0
            else:
                sim = cosine_similarity(query_emb, chunk_emb)
            scored.append((match_item, sim))

        scored.sort(key=lambda x: x[1], reverse=True)
        top = scored[:top_k]
        return [x[0]["text"] for x in top]
    else:
        # No exact matches found => fallback to BM25 if requested
        if not use_bm25_fallback:
            return _embed_retrieve(query, chunk_data, top_k)
        else:
            # BM25 approach
            # 1) Build a corpus: each chunk => token list
            corpus = []
            for item in chunk_data:
                if hasattr(item["text"], "page_content"):
                    txt = item["text"].page_content
                else:
                    txt = str(item["text"])
                corpus.append(txt.split())

            bm25 = BM25Okapi(corpus)
            tokenized_query = query.split()
            scores = bm25.get_scores(tokenized_query)

            # Pair each chunk with its BM25 score
            chunk_scores = list(zip(chunk_data, scores))
            chunk_scores.sort(key=lambda x: x[1], reverse=True)
            top_bm25 = chunk_scores[:max(10, top_k)]  # keep more, then re-rank

            # Re-rank by embeddings if you like
            query_emb = _embed_query(query)
            re_scored = []
            for (ch_item, bm25_score) in top_bm25:
                chunk_emb = ch_item["embedding"]
                if not chunk_emb:
                    sim = 0.0
                else:
                    sim = cosine_similarity(query_emb, chunk_emb)
                # Combine or ignore BM25? Example: final_score = sim + (bm25_score/100)
                # For simplicity, let's sum them:
                final_score = sim + (bm25_score / 100.0)
                re_scored.append((ch_item, final_score))

            re_scored.sort(key=lambda x: x[1], reverse=True)
            final = re_scored[:top_k]
            return [x[0]["text"] for x in final]

def _embed_query(query_str: str) -> List[float]:
    """
    Helper that embeds a search query ONCE. If you want to skip all query embeddings
    altogether, you can remove embedding-based re-rank or keep a cache.
    """
    try:
        print("The text being sent is :", query_str)
        response = openai.embeddings.create(
            model="text-embedding-3-large",
            input=query_str
        )
        # 'data' is a list of embeddings; we only sent one input, so it's 0th
        return response.data[0].embedding
    except Exception as e:
        print(f"[Warning] Query embedding failed => {e}")
        return []

def _embed_retrieve(query: str, chunk_data: List[Dict[str,Any]], top_k: int) -> List[Any]:
    """
    Fallback method: do a pure embedding-based retrieval across all chunks
    (use precomputed embeddings, skip any chunk that has no embedding).
    """
    query_emb = _embed_query(query)
    scored = []
    for item in chunk_data:
        c_emb = item["embedding"]
        if not c_emb:
            sim = 0.0
        else:
            sim = cosine_similarity(query_emb, c_emb)
        scored.append((item, sim))

    scored.sort(key=lambda x: x[1], reverse=True)
    final = scored[:top_k]
    return [x[0]["text"] for x in final]

def fallback_data(name: str) -> dict:
    """
    Return a fallback result for an ingredient when no information could be found.
    """
    return {
        "Name": name,
        "Synonyms": None,
        "StandardName": None,
        "FunctionalRole": None,
        "Citation": None
    }


def precompute_embeddings(chunks, concurrency=5):
    """
    Precompute embeddings for a list of chunks using concurrency.
    
    :param chunks: The list of chunk documents/strings.
    :param concurrency: Number of concurrent workers to use.
    :return: A list of dictionaries, each containing {'text': original chunk, 'embedding': [...]}.
    """
    def process_chunk(doc):
        text_str = doc.page_content if hasattr(doc, "page_content") else str(doc)
        chunk_emb = embed_text(text_str)  # Calls your embed_text function
        return {
            "text": doc,
            "embedding": chunk_emb
        }

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        # Map each chunk to our process_chunk function
        results = list(executor.map(process_chunk, chunks))

    return results


import json
import csv
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, Any, List


#### Ist Refinement Cycle ( LLM Enabled )
def process_all_ingredients_simple(
    unique_ing_dict,
    chunk_data,
    output_docx="refined_ingredients.docx",
    output_csv="refined_ingredients.csv",
    take_user_input=True
):
    """
    1) For each ingredient in unique_ing_dict, find any relevant chunks.
    2) If none found, keep original data.
    3) If found, call refine_ingredient_data_via_llm function:
       - If take_user_input=True => do sequential calls with user prompts.
       - If take_user_input=False => do concurrency with batches of 10 items (no user prompt).
    4) Write the final data (either refined or original) into:
       - A Word table with columns:
         [NAME, PLACEHOLDER, ROLE_OF_PLACEHOLDER, STANDARD_NAMES, SOURCE_SUPPLIER,
          SUPPLYING_COMPANY, FUNCTIONAL_ROLES, MENTIONED_EXAMPLES,
          IS_ACTIVE_INGREDIENT, IS_STRUCTURAL_COMPONENT]
       - A CSV file with the same columns
    """

    from concurrent.futures import ThreadPoolExecutor, as_completed

    # We'll store data to refine in a list, then run concurrency if needed.
    # final_results: list of dicts with keys:
    #   "ing_name", "refined_data" (or None)
    #   "no_chunks" (bool), "original_data_for_llm" (dict)
    concurrency_batch = []
    final_results = []

    # Prepare all items
    for ing_name, ing_info in unique_ing_dict.items():
        # Ensure dict
        if not isinstance(ing_info, dict):
            if isinstance(ing_info, str):
                try:
                    parsed = json.loads(ing_info)
                    ing_info = parsed if isinstance(parsed, dict) else {}
                except json.JSONDecodeError:
                    ing_info = {}
            else:
                ing_info = {}

        # Provide defaults
        ing_info.setdefault("standard_names", [])
        ing_info.setdefault("functional_roles", [])
        ing_info.setdefault("is_active_ingredient", False)
        ing_info.setdefault("is_structural_component", False)
        ing_info.setdefault("examples_mentioned", [])

        # Attempt to find relevant chunks
        relevant_chunks = find_relevant_chunks(ing_name, chunk_data, top_k=7)

        original_data_for_llm = {"original_ingredient_name": ing_name}
        for k, v in ing_info.items():
            original_data_for_llm[f"original_{k}"] = v

        if not relevant_chunks:
            # No chunks => no refinement
            final_results.append({
                "ing_name": ing_name,
                "no_chunks": True,
                "refined_data": {},
                "original_data_for_llm": original_data_for_llm
            })
        else:
            # We have chunks
            if take_user_input:
                # No concurrency => just call refine_ingredient_data_via_llm immediately
                refined_data = refine_ingredient_data_via_llm(
                    original_data_for_llm,
                    relevant_chunks,
                    take_user_input=True
                )
                final_results.append({
                    "ing_name": ing_name,
                    "no_chunks": False,
                    "refined_data": refined_data,
                    "original_data_for_llm": original_data_for_llm
                })
            else:
                # concurrency => queue it up for batch
                concurrency_batch.append((
                    ing_name,
                    original_data_for_llm,
                    relevant_chunks
                ))

    # Now handle concurrency if take_user_input=False
    # We'll process concurrency in BATCHES of 10 items
    def _concurrent_call(params):
        ing_name, orig_data, chks = params
        # We'll call refine_ingredient_data_via_llm with take_user_input=False
        # so it auto-accepts the LLM result
        return ing_name, refine_ingredient_data_via_llm(
            orig_data,
            chks,
            take_user_input=False
        )

    if not take_user_input and concurrency_batch:
        # We run concurrency in chunks of 10
        batch_size = 10
        i_start = 0
        while i_start < len(concurrency_batch):
            sub_batch = concurrency_batch[i_start: i_start+batch_size]
            i_start += batch_size
            # Launch them
            results_dict = {}
            with ThreadPoolExecutor(max_workers=10) as executor:
                futures = {executor.submit(_concurrent_call, sb): sb[0] for sb in sub_batch}
                for fut in as_completed(futures):
                    ing_name = futures[fut]
                    try:
                        name, data = fut.result()
                        results_dict[name] = data
                    except Exception as e:
                        print(f"[Error] concurrency for {ing_name} => {e}")
                        results_dict[ing_name] = {}

            # Merge concurrency results into final_results
            for (ing_name, orig_data, chks) in sub_batch:
                final_results.append({
                    "ing_name": ing_name,
                    "no_chunks": False,
                    "refined_data": results_dict.get(ing_name, {}),
                    "original_data_for_llm": orig_data
                })

    # --------------- Now write results to doc/csv ---------------

    doc = Document()
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text = "NAME"
    hdr[1].text = "PLACEHOLDER"
    hdr[2].text = "ROLE_OF_PLACEHOLDER"
    hdr[3].text = "STANDARD_NAMES"
    hdr[4].text = "SOURCE_SUPPLIER"
    hdr[5].text = "SUPPLYING_COMPANY"
    hdr[6].text = "FUNCTIONAL_ROLES"
    hdr[7].text = "MENTIONED_EXAMPLES"
    hdr[8].text = "IS_ACTIVE_INGREDIENT"
    hdr[9].text = "IS_STRUCTURAL_COMPONENT"

    with open(output_csv, mode="w", newline="", encoding="utf-8") as csvfile:
        csv_writer = csv.writer(csvfile)
        csv_writer.writerow([
            "NAME",
            "PLACEHOLDER",
            "ROLE_OF_PLACEHOLDER",
            "STANDARD_NAMES",
            "SOURCE_SUPPLIER",
            "SUPPLYING_COMPANY",
            "FUNCTIONAL_ROLES",
            "MENTIONED_EXAMPLES",
            "IS_ACTIVE_INGREDIENT",
            "IS_STRUCTURAL_COMPONENT"
        ])

        # ① keep a list of rows so we can build a DataFrame later
        df_rows = []

        # Process final results
        for item in final_results:
            ing_name = item["ing_name"]
            refined_data = item["refined_data"]
            original_data = item["original_data_for_llm"]

            if refined_data:
                name_val = refined_data.get("LLM_refined_name", "")
                placeholder = refined_data.get("LLM_refined_name_placeholder_for", [])
                role_placeholder = refined_data.get("LLM_refined_role_of_placeholder", [])
                std_names_val = refined_data.get("LLM_refined_standard_names", [])
                source_val = refined_data.get("LLM_refined_source_supplier", [])
                supply_co_val = refined_data.get("LLM_refined_supplying_company", [])
                func_roles_val = refined_data.get("LLM_refined_functional_roles", [])
                ex_val = refined_data.get("LLM_refined_mentioned_examples", [])
                is_active_ingredient = refined_data.get("LLM_refined_is_active_ingredient", False)
                is_structural_component = refined_data.get("LLM_refined_is_structural_ingredient", False)
            else:
                # Fall back to original data
                name_val = ing_name
                placeholder = []
                role_placeholder = []
                std_names_val = original_data.get("original_standard_names", [])
                source_val = []
                supply_co_val = []
                func_roles_val = original_data.get("original_functional_roles", [])
                ex_val = original_data.get("original_examples_mentioned", [])
                is_active_ingredient = original_data.get("original_is_active_ingredient", False)
                is_structural_component = original_data.get("original_is_structural_component", False)

            # Convert lists
            placeholder_str = ", ".join(placeholder) if isinstance(placeholder, list) else str(placeholder)
            role_placeholder_str = ", ".join(role_placeholder) if isinstance(role_placeholder, list) else str(role_placeholder)
            std_names_str = ", ".join(std_names_val) if isinstance(std_names_val, list) else str(std_names_val)
            source_str = ", ".join(source_val) if isinstance(source_val, list) else str(source_val)
            supply_str = ", ".join(supply_co_val) if isinstance(supply_co_val, list) else str(supply_co_val)
            func_roles_str = ", ".join(func_roles_val) if isinstance(func_roles_val, list) else str(func_roles_val)
            ex_str = ", ".join(ex_val) if isinstance(ex_val, list) else str(ex_val)
            is_active_str = "True" if is_active_ingredient else "False"
            is_structural_str = "True" if is_structural_component else "False"

            # Word table row
            row = table.add_row().cells
            row[0].text = str(name_val)
            row[1].text = placeholder_str
            row[2].text = role_placeholder_str
            row[3].text = std_names_str
            row[4].text = source_str
            row[5].text = supply_str
            row[6].text = func_roles_str
            row[7].text = ex_str
            row[8].text = is_active_str
            row[9].text = is_structural_str

            # CSV row
            csv_writer.writerow([
                str(name_val),
                placeholder_str,
                role_placeholder_str,
                std_names_str,
                source_str,
                supply_str,
                func_roles_str,
                ex_str,
                is_active_str,
                is_structural_str
            ])

            # -------- add the same record to df_rows ----------
            df_rows.append({
                "NAME": name_val,
                "PLACEHOLDER": placeholder_str,
                "ROLE_OF_PLACEHOLDER": role_placeholder_str,
                "STANDARD_NAMES": std_names_str,
                "SOURCE_SUPPLIER": source_str,
                "SUPPLYING_COMPANY": supply_str,
                "FUNCTIONAL_ROLES": func_roles_str,
                "MENTIONED_EXAMPLES": ex_str,
                "IS_ACTIVE_INGREDIENT": is_active_str,
                "IS_STRUCTURAL_COMPONENT": is_structural_str,
            })

    doc.save(output_docx)
    print(f"Document saved as '{output_docx}'")
    print(f"CSV file saved as '{output_csv}'")

    # ②  Build DataFrame and return it together with the file paths
    import pandas as pd
    refined_df = pd.DataFrame(df_rows) 

    # NEW – hand paths back to the caller
    return refined_df, os.path.abspath(output_docx), os.path.abspath(output_csv)


import json
import re
import pandas as pd
from typing import Dict, Any, List
import openai
from concurrent.futures import ThreadPoolExecutor, as_completed

def refine_deduplicated_dataframe(
    input_csv: pd.DataFrame,
    chunk_data,  # your existing data structure used for "relevant chunks"
    include_relevant_chunks: bool = True,
    output_csv: str = "WO2023060387-CORE-ACTIVES.csv",
    openai_model: str = "gpt-4o",
    take_user_input: bool = False,
    max_workers: int = 10
):
    """
    1) Loads a 'deduplicated' CSV file (e.g., from union-find merges).
    2) If take_user_input=False => uses concurrency with up to `max_workers` calls in parallel.
       If take_user_input=True => processes each row sequentially, prompting user acceptance or discard.
    3) For each row:
       - Optionally includes relevant chunk text for row['NAME'] (if include_relevant_chunks=True).
       - Asks the LLM to produce a refined JSON structure that may split
         an overstuffed row into multiple smaller entries.
       - Each new entry must have "representative_member" plus the usual fields.
    4) We parse the returned JSON, collect all new “mini-rows,” and output them to `output_csv`.
    
    The final CSV includes these columns (one row per “mini-entry”):
      - "representative_member"
      - "split_entries"       (the original or refined name)
      - "standard_names"
      - "suppliers"
      - "functional_roles"
      - "examples"
      - "is_active"
      - "is_structural"
      - "relation_to_representative"
    
    If a row is actually single, it remains a single entry. If it's multi, it may be split into multiple lines.

    :param input_csv: path to CSV from the deduplication step
    :param chunk_data: your chunk data to find relevant context
    :param include_relevant_chunks: if True, we find & pass chunk text to the LLM
    :param output_csv: where to save final refined data
    :param openai_model: which OpenAI model to use, e.g. 'gpt-4'
    :param take_user_input: if True => no concurrency, user is prompted to accept or discard each row's result
    :param max_workers: concurrency level if take_user_input=False
    """

    # 1) Load the CSV into a dataframe
    # df = pd.read_csv(input_csv).fillna("")
    df = input_csv.copy().fillna("")
    all_rows = df.to_dict(orient="records")

    # Helper to find relevant chunks
    def get_relevant_chunk_text(row: dict) -> str:
        """ If include_relevant_chunks, find chunk text for row['NAME']; else return empty string. """
        name_val = row.get("NAME", "")
        if not name_val:
            return "(No name field to find relevant chunks.)"
        # You must have a find_relevant_chunks function that returns text
        relevant_text = find_relevant_chunks(name_val, chunk_data, top_k=5)
        if not relevant_text:
            return "(No relevant chunks found.)"
        return relevant_text

    def clean_llm_response(raw_text: str) -> str:
        """
        Remove code-fence artifacts (```json ... ```),
        any leading "json\n" or "json ",
        and then strip whitespace. 
        This helps ensure valid JSON before parsing.
        """
        # Remove the triple backtick code fences if present
        cleaned = raw_text.replace("```json", "").replace("```", "")
        # Remove any leading "json\n" or "json "
        cleaned = re.sub(r'^json\s*', '', cleaned.strip(), flags=re.IGNORECASE)
        return cleaned.strip()

    def refine_one_row(row: dict, model: str, interactive: bool) -> List[dict]:
        """
        1) Optionally includes relevant chunk text for the row.
        2) Sends a prompt to the LLM to split or refine the row.
        3) Parses the JSON response and maps it into the final schema (dicts).
        4) Returns a list of dicts (each “mini-entry”).
        """
        # Possibly fetch chunk text
        chunk_text = get_relevant_chunk_text(row) if include_relevant_chunks else ""
        row_json = json.dumps(row, ensure_ascii=False)

        # System instructions
        system_instructions = (
            "You are an assistant that refines or splits deduplicated cosmetic-ingredient rows.\n"
            "Return ONLY valid JSON with an array of refined entries. Each new entry must have:\n"
            '- "LLM_refined_representative_member"\n'
            '- "LLM_refined_name"\n'
            '- "LLM_refined_cluster_relation"\n'
            '- "LLM_refined_standard_names"\n'
            '- "LLM_refined_source_supplier"\n'
            '- "LLM_refined_supplying_company"\n'
            '- "LLM_refined_functional_roles"\n'
            '- "LLM_refined_mentioned_examples"\n'
            '- "LLM_refined_is_active_ingredient"\n'
            '- "LLM_refined_is_structural_ingredient"\n'
            "If a row is truly single, keep it one entry. If multiple subsets, produce multiple entries.\n"
            "If something is unknown, keep it empty.\n"
        )

        
        # The user prompt
        user_prompt = f"""You have 1 deduplicated row in JSON form, plus optional chunk text:
        Original Row Data:
        {row_json}
        
        Your tasks are:
        0. MOST IMPORTANT THERE SHOULD BE NO REPEATED ENTRIES IF TWO ENTRIES ARE JUST DIFFERENT IN LOWER CASE UPPER CASE AND HAVE NO DIFFERENCES , 
        REMOVE THE REDUNDANT ENTRY , NOT TO BE APPLIED ON INDIVIDUAL FIELDS , ONLY IF TWO ENTRIES ARE EXACTLY SAME IN  A ROW AND REPEATED YOU CAN REMOVE
        ONE.
        1. Analyze each row/cluster and determine the most representative member — the chemical name that best describes the cluster.
        2. Each row is a combination of similar ingredients or representation , I dont want to loose that base information , but still want to 
        group it properly , right now the one row represents a group of similar names , I want to you create entries in such that 
        we can capture hierarchial relationships between the individual names and the representative member so if a group is example - degradable 
        implants based on collagen, hyaluronic acid, or polylactic acid, resorbable components ,and you think that xyz is a good representative member
        you will give me properly broken entities with their hierarchical or any kind of relation to the representative member ,so in the worst case 
        this individual example row will have 3 different entities , each having a common representative member and their own relation to this 
        representative member , example in the worst case this one row will create 4 json entries , because each has its own relationship to the 
        representative member , but in some cases you might find that two names are very equal and they dont share a unique relationship with 
        the representative member hence you can merge those entries together , but the relationship of the entities split should be unique to the 
        chosen representative member.
        3. If a cluster contains multiple candidate representative members in other words it might be overstuffed and might contains more than one cluster
        information in which case first you need, split the cluster accordingly into separate groups , chose representative members for each group 
        and then do step 2 on each group to get the correct entities , so please note we are not loosing data any point we are just rearranging it 
        based on the representative member so that we can actually create clusters.
        4. For each refined entry, return the output strictly as valid JSON with the following keys:
           - "LLM_refined_name": The original concatenated name from the row provided chemical name.
           - "LLM_refined_representative_member": The chosen representative member name for the cluster.
           - "LLM_refined_cluster_relation": A brief description of how the names in the cluster relate to the representative member.
           - "LLM_refined_standard_names": A list of standard names, based in the name splits its standard names accordingly from the names made available to you.
           - "LLM_refined_source_supplier": A list of source suppliers based in the name splits its standard names accordingly from the names made available to you.
           - "LLM_refined_supplying_company": A list of supplying companies.
           - "LLM_refined_functional_roles": A list of functional roles.
           - "LLM_refined_mentioned_examples": A list of mentioned examples.
           - "LLM_refined_is_active_ingredient": Boolean flag indicating if it is an active ingredient.
           - "LLM_refined_is_structural_ingredient": Boolean flag indicating if it is a structural component.
           
        5. Consider the following relevant context in your analysis:
        {chunk_text}
        6. Return the output as a list of refined entries in JSON format. Example output format :
          {{
          "LLM_refined_name": The specific name of the ingredient as given in the row , try not to change this example - 
          degradable implants based on collagen, hyaluronic acid, or polylactic acid, resorbable components , in this collagen is one 
          refined name it belongs to "degradable implants" representative member becuase Collagen is used as a primary component in degradable implants, 
          often combined with other resorbable materials, so degradable implants is a parent of collagen , use this semantic to figure out the 
          refined name and representative member for all the rows , the refined name is exactly taken from the original row , but the representative member
          actual informs the non-technical user the role of the refined name in the bigger picture.
          "LLM_refined_representative_member": Representative Name  that describes the relation from the representative to the refined name
          "LLM_refined_cluster_relation": "Describes how the refined name of the cluster relate to this representative member chosen for the cluster"
          "LLM_refined_standard_names": [
            "Standard Name A",
            "Standard Name B"
          ],
          "LLM_refined_source_supplier": [
            "Supplier A",
            "Supplier B"
          ],
          "LLM_refined_supplying_company": [
            "Company A",
            "Company B"
          ],
          "LLM_refined_functional_roles": [
            "Role A",
            "Role B"
          ],
          "LLM_refined_mentioned_examples": [
            "Example A",
            "Example B"
          ],
          "LLM_refined_is_active_ingredient": false,
          "LLM_refined_is_structural_ingredient": true,
          
        }}

        """

        # Step D: Call the LLM
        try:
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_instructions},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3
            )
            llm_text = response.choices[0].message.content
           
        except Exception as e:
            print(f"[Error] refine_ingredient_data_via_llm => {e}")
            return []
        # Clean up the returned text (remove code fences, "json\n", etc.)
        llm_text_clean = clean_llm_response(llm_text)

        # Attempt to parse the JSON
        refined_entries = []
        try:
            parsed = json.loads(llm_text_clean)
            # If only a single dict was returned, wrap it in a list
            if isinstance(parsed, dict):
                parsed = [parsed]
        except json.JSONDecodeError:
            print("[Warning] LLM did not return valid JSON. Full text:")
            print(llm_text_clean)
            return []

        # Now convert each parsed item to the final schema for CSV
        for ent in parsed:
            # We map the LLM keys to your final CSV columns
            new_entry = {
                "representative_member": ent.get("LLM_refined_representative_member", ""),
                "split_entries": ent.get("LLM_refined_name", ""),
                "standard_names": ent.get("LLM_refined_standard_names", []),
                "suppliers": ent.get("LLM_refined_source_supplier", []) 
                             + ent.get("LLM_refined_supplying_company", []),
                "functional_roles": ent.get("LLM_refined_functional_roles", []),
                "examples": ent.get("LLM_refined_mentioned_examples", []),
                "is_active": ent.get("LLM_refined_is_active_ingredient", False),
                "is_structural": ent.get("LLM_refined_is_structural_ingredient", False),
                "relation_to_representative": ent.get("LLM_refined_cluster_relation", "")
            }
            refined_entries.append(new_entry)

        # If in interactive mode, show the user and let them accept or discard
        if interactive:
            if refined_entries:
                print("\nRefined JSON from LLM:\n", json.dumps(refined_entries, indent=2, ensure_ascii=False))
                user_decision = input("\nPress ENTER to accept, or 'n' to discard: ").strip().lower()
                if user_decision == 'n':
                    print("[Info] user discarded.")
                    return []
                else:
                    print("[Info] user accepted.")
                    return refined_entries
            else:
                print("[Info] LLM returned empty; nothing to accept => returning empty.")
                return []
        else:
            # Automatic acceptance
            return refined_entries

    # List of final sub-entries from all rows
    final_entries: List[dict] = []

    # 2) run in either interactive (sequential) or concurrent mode
    if take_user_input:
        # SEQUENTIAL
        for i, row_data in enumerate(all_rows):
            print(f"\n[Processing row {i+1}/{len(all_rows)} in sequential mode]")
            sub_entries = refine_one_row(row_data, openai_model, interactive=True)
            final_entries.extend(sub_entries)
    else:
        # CONCURRENT
        def concurrent_worker(idx: int, row_data: dict):
            print(f"[Concurrent] Starting row {idx+1}/{len(all_rows)} ...")
            return refine_one_row(row_data, openai_model, interactive=False)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {}
            for i, row_data in enumerate(all_rows):
                fut = executor.submit(concurrent_worker, i, row_data)
                futures[fut] = i

            for fut in as_completed(futures):
                row_idx = futures[fut]
                try:
                    result = fut.result()
                    if result: 
                        final_entries.extend(result)
                except Exception as e:
                    print(f"[Error] concurrency row {row_idx} => {e}")

    # 3) Convert final_entries to a DataFrame with your desired columns
    out_df = pd.DataFrame(
        final_entries,
        columns=[
            "representative_member",
            "split_entries",
            "standard_names",
            "suppliers",
            "functional_roles",
            "examples",
            "is_active",
            "is_structural",
            "relation_to_representative"
        ]
    )

    # Ensure all columns exist
    for col in [
        "representative_member",
        "split_entries",
        "standard_names",
        "suppliers",
        "functional_roles",
        "examples",
        "is_active",
        "is_structural",
        "relation_to_representative"
    ]:
        if col not in out_df.columns:
            out_df[col] = ""

    # Write CSV
    out_df.to_csv(output_csv, index=False)

    print(f"\n[Done] Processed {len(all_rows)} original rows => {len(out_df)} refined entries => saved '{output_csv}'.")
    return out_df


##### ALL THE GRAPH BASED OUTPUTS - 
###############################################################################
# MONTE-CARLO MULTI-THRESHOLD DEDUPLICATION WITH OPENAI EMBEDDINGS
# + Additional Graphs for Unique & Repeated Counting, docx & CSV export
#
# EXPLANATION OF SIMILARITY THRESHOLD:
#   - The threshold is a value from 0.0 to 1.0 that decides how similar two rows'
#     embeddings must be before we consider them "duplicates" and merge them.
#   - For example:
#       threshold = 0.00  => merge everything that has a similarity >= 0.00
#                            (basically merges all rows into one big cluster).
#       threshold = 0.50  => merges only pairs of rows whose embedding similarity
#                            is at least 50%. 
#       threshold = 1.00  => merges only pairs of rows that are 100% embedding-
#                            identical (extremely strict).
#   - As the threshold increases (0→1):
#       * We allow fewer merges (they must be more similar to merge).
#       * Usually the number of "clusters" (rows in the fused DataFrame) grows.
#
# WHAT "NO REPEATED ENTRIES" MEANS:
#   - After deduplication at a given threshold, we look at the resulting DataFrame.
#   - We check if any two rows have the same NAME (case-insensitive).
#   - If none are the same (ignoring case), we say "no repeated entries."
#
# TIERED CRITERIA:
#   1) DEEP-KNOWLEDGE: "No repeated entries" in the final DataFrame. 
#   2) MEDIUM-SCALE:   "No repeated entries" AND the final DF has <= 50 rows.
#   3) CORE-LEVEL:     "No repeated entries" AND the final DF has <= 9 rows.
#
# OUTPUTS:
#   - For each threshold, a CSV of the deduplicated data: "output_at_{thr:.2f}.csv"
#   - Four plots (PNG):
#       1) row_count_vs_threshold.png        (clusters vs. threshold)
#       2) unique_count_vs_threshold.png     (unique NAME count vs. threshold)
#       3) repeated_count_vs_threshold.png   (repeated NAME count vs. threshold)
#       4) combined_plot.png                 (all lines + vertical lines for
#                                             chosen thresholds)
#   - A CSV "thresholds_summary.csv" with columns:
#        threshold, row_count, unique_count, repeated_count, no_repeats_bool
#   - A Word document "thresholds_summary.docx" containing the same summary data.
###############################################################################


###############################################################################
# MONTE-CARLO MULTI-THRESHOLD DEDUPLICATION WITH OPENAI EMBEDDINGS
# + Additional Graphs, docx & CSV export
#
# NEW LOGIC (2023-xx-xx):
#   1) "Full-Analysis" threshold:
#       - The deduplicated DataFrame must have no repeated entries.
#       - Among those, pick the threshold with the MAX unique_count.
#         If there's a tie in unique_count, pick the HIGHEST threshold.
#   2) "Core-Level" threshold:
#       - No repeated entries AND row_count <= 9.
#       - If multiple meet that, pick the HIGHEST threshold.
#   3) Everything else => "Partial" (no vertical line on the plot).
#
#   We remove the old "medium scale" logic. If you want it, you can mention
#   it in the docx text, but we don't compute or highlight it in code.
#
#   The Combined Plot now has exactly two vertical lines (if they exist).
#   - If both the "full analysis" & "core-level" thresholds are the same,
#     we draw ONE line labeled "Full + Core thr=??".
#
#   The x-axis now has ticks every 0.1 from 0.0 to 1.0, with minor ticks on.
#   If row_count == unique_count at some threshold, we add a note/annotation.
###############################################################################

import os
import math
import numpy as np
import pandas as pd
import openai
from typing import List
import matplotlib.pyplot as plt

# For Word output:
from docx import Document
from docx.shared import Inches

############################################################################
# 0) SETUP: Provide your OpenAI API key. (Replace or use environment variable.)
############################################################################
openai.api_key = openai_key
DEFAULT_EMBED_MODEL = "text-embedding-3-large"

############################################################################
# 1) COSINE SIMILARITY
############################################################################
def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    dot = sum(a * b for a, b in zip(vec1, vec2))
    mag1 = math.sqrt(sum(a*a for a in vec1))
    mag2 = math.sqrt(sum(a*a for a in vec2))
    if not mag1 or not mag2:
        return 0.0
    return dot / (mag1 * mag2)

############################################################################
# 2) GET OPENAI EMBEDDING
############################################################################
def get_embedding(text: str, model: str = DEFAULT_EMBED_MODEL) -> List[float]:
    try:
        print(f"[DEBUG] Embedding text: {text}")
        response = openai.embeddings.create(model=model, input=text)
        return response.data[0].embedding
    except Exception as e:
        print(f"[Error] Embedding failed for '{text[:30]}': {e}")
        return []

############################################################################
# 3) UNION-FIND
############################################################################
class UnionFind:
    def __init__(self, n: int):
        self.parent = list(range(n))
        self.rank = [0]*n

    def find(self, x: int) -> int:
        if self.parent[x] != x:
            self.parent[x] = self.find(self.parent[x])
        return self.parent[x]

    def union(self, x: int, y: int):
        rx = self.find(x)
        ry = self.find(y)
        if rx != ry:
            if self.rank[rx] < self.rank[ry]:
                self.parent[rx] = ry
            elif self.rank[rx] > self.rank[ry]:
                self.parent[ry] = rx
            else:
                self.parent[ry] = rx
                self.rank[rx] += 1

############################################################################
# 4) FUSE ROWS => SINGLE ROW
############################################################################
def fuse_rows(df: pd.DataFrame, indices: List[int]) -> dict:
    """
    Merge all columns from a set of rows by set-union of their values
    (comma-based for strings).
    """
    if not indices:
        return {}
    fused = dict(df.iloc[indices[0]])
    for idx in indices[1:]:
        row = df.iloc[idx]
        for col in df.columns:
            old_val = fused[col]
            new_val = row[col]
            if pd.isna(old_val) and pd.isna(new_val):
                continue
            set_old = _value_to_set(old_val)
            set_new = _value_to_set(new_val)
            merged = set_old.union(set_new)
            if len(merged) == 0:
                fused[col] = np.nan
            elif len(merged) == 1:
                fused[col] = list(merged)[0]
            else:
                fused[col] = ", ".join(sorted(merged))
    return fused

def _value_to_set(val) -> set:
    if pd.isna(val):
        return set()
    if isinstance(val, str):
        parts = [x.strip() for x in val.split(",") if x.strip()]
        return set(parts)
    if isinstance(val, list):
        return set(str(x).strip() for x in val)
    return {str(val).strip()}

############################################################################
# 5) DEDUPLICATION PASS
############################################################################
def deduplicate_one_pass(df: pd.DataFrame, similarity_matrix: List[List[float]], threshold: float) -> pd.DataFrame:
    """
    Single-pass union-find clustering for similarity >= threshold.
    """
    n = len(df)
    uf = UnionFind(n)
    for i in range(n):
        for j in range(i+1, n):
            if similarity_matrix[i][j] >= threshold:
                uf.union(i, j)

    # Build clusters
    clusters = {}
    for i in range(n):
        root = uf.find(i)
        if root not in clusters:
            clusters[root] = []
        clusters[root].append(i)

    # Fuse each cluster
    fused_rows = []
    for root_idx, members in clusters.items():
        merged_dict = fuse_rows(df, members)
        fused_rows.append(merged_dict)

    return pd.DataFrame(fused_rows).reset_index(drop=True)

############################################################################
# 6) CASE-INSENSITIVE REPEATED NAME COUNT
############################################################################
def count_repeated_names(df: pd.DataFrame, name_col: str) -> (int, int):
    """
    Return (unique_count, repeated_count).
      unique_count = # of distinct names ignoring case
      repeated_count = how many total rows share a name that appears >1 times
    """
    lower_names = df[name_col].fillna("").astype(str).str.strip().str.lower()
    freq_map = lower_names.value_counts()
    unique_count = len(freq_map)
    repeated_rows = 0
    for nm, freq in freq_map.items():
        if freq > 1:
            repeated_rows += freq
    return (unique_count, repeated_rows)

############################################################################
# 7) MONTE CARLO ANALYSIS
############################################################################
def monte_carlo_analysis(
    df: pd.DataFrame,
    output_dir: str,
    name_col: str = "NAME",
    model_name: str = DEFAULT_EMBED_MODEL,
    threshold_step: float = 0.05
):
    """
    Performs the multi-threshold dedup, removes repeated entries if similarity >= threshold.
    Then:
      - Saves CSV for each threshold
      - Collects row_count, unique_count, repeated_count, no_repeats_bool
      - Picks "Full Analysis" threshold (max unique_count, no repeats; tie->highest thr)
      - Picks "Core-Level" threshold (no repeats, row_count <= 9; tie->highest thr)
      - Plots row_count, unique_count, repeated_count vs. threshold
      - On combined plot, draws short vertical lines from y=0 up to max of the three 
        metrics at that threshold. No "row=unique" text annotations.
      - Saves summary CSV & summary docx.
    """
    os.makedirs(output_dir, exist_ok=True)

    ################################################################
    # A) Embeddings
    ################################################################
    df = df.copy().reset_index(drop=True)
    names = df[name_col].fillna("").astype(str)
    print(f"[INFO] Embedding {len(names)} rows with model={model_name} ...")
    all_embeddings = []
    for nm in names:
        emb = get_embedding(nm, model=model_name)
        all_embeddings.append(emb)

    ################################################################
    # B) NxN similarity
    ################################################################
    n = len(df)
    print("[INFO] Building NxN similarity matrix ...")
    sim_matrix = [[0.0]*n for _ in range(n)]
    for i in range(n):
        for j in range(i+1, n):
            if all_embeddings[i] and all_embeddings[j]:
                sim = cosine_similarity(all_embeddings[i], all_embeddings[j])
                sim_matrix[i][j] = sim
                sim_matrix[j][i] = sim

    ################################################################
    # C) Threshold loop
    ################################################################
    thresholds = np.arange(0.0, 1.0001, threshold_step)
    thresholds = [round(t,2) for t in thresholds if t <= 1.0]

    results = []
    thr2df    = {}          # <── NEW
    for thr in thresholds:
        ddf = deduplicate_one_pass(df, sim_matrix, thr)
        # keep a copy for later retrieval
        thr_key = f"{thr:.2f}"  # always stored with the SAME string format
        thr2df[thr_key] = ddf      # <── NEW
        row_count = len(ddf)
        unique_cnt, repeated_cnt = count_repeated_names(ddf, name_col)
        no_repeats_bool = (repeated_cnt == 0)

        # Save CSV
        out_name = f"output_at_{thr_key}.csv"
        out_path = os.path.join(output_dir, out_name)
        ddf.to_csv(out_path, index=False)

        row_info = {
            "threshold": thr,
            "row_count": row_count,
            "unique_count": unique_cnt,
            "repeated_count": repeated_cnt,
            "no_repeats_bool": no_repeats_bool
        }
        results.append(row_info)
        print(f"[THR={thr_key}] row_count={row_count}, unique={unique_cnt}, repeated={repeated_cnt}, no_repeats={no_repeats_bool}")

    summary_df = pd.DataFrame(results)

    ################################################################
    # D) Determine Full-Analysis & Core-Level
    ################################################################
    # Full Analysis => no repeats, max unique_count, tie->highest thr
    fa_candidates = summary_df[summary_df["no_repeats_bool"] == True]
    full_analysis_threshold = float("nan")
    if not fa_candidates.empty:
        max_uc = fa_candidates["unique_count"].max()
        subset = fa_candidates[fa_candidates["unique_count"] == max_uc]
        full_analysis_threshold = subset["threshold"].max()  # highest threshold

    # Core-Level => no repeats & row_count <= 9, tie->highest thr
    cl_candidates = fa_candidates[fa_candidates["row_count"] <= 9]
    core_level_threshold = float("nan")
    if not cl_candidates.empty:
        core_level_threshold = cl_candidates["threshold"].max()


    # Convert the numeric answers to the same **string** key we used above
    full_thr_key = (f"{full_analysis_threshold:.2f}"
                    if not np.isnan(full_analysis_threshold) else None)
    core_thr_key = (f"{core_level_threshold:.2f}"
                    if not np.isnan(core_level_threshold) else None)

    ################################################################
    # E) Individual Plots
    ################################################################
    # row_count
    plt.figure(figsize=(7,5))
    plt.plot(summary_df["threshold"], summary_df["row_count"], marker='o', label="Row Count")
    plt.xlabel("Threshold")
    plt.ylabel("Row (Cluster) Count")
    plt.title("Row Count vs. Threshold")
    plt.grid(True)
    plt.xticks(np.arange(0,1.01,0.1))
    plt.minorticks_on()
    plt.savefig(os.path.join(output_dir, "row_count_vs_threshold.png"), dpi=150)
    plt.close()

    # unique_count
    plt.figure(figsize=(7,5))
    plt.plot(summary_df["threshold"], summary_df["unique_count"], marker='o', color="green", label="Unique NAMEs")
    plt.xlabel("Threshold")
    plt.ylabel("Unique NAMEs")
    plt.title("Unique NAME count vs. Threshold")
    plt.grid(True)
    plt.xticks(np.arange(0,1.01,0.1))
    plt.minorticks_on()
    plt.savefig(os.path.join(output_dir, "unique_count_vs_threshold.png"), dpi=150)
    plt.close()

    # repeated_count
    plt.figure(figsize=(7,5))
    plt.plot(summary_df["threshold"], summary_df["repeated_count"], marker='o', color="red", label="Repeated NAMEs")
    plt.xlabel("Threshold")
    plt.ylabel("Repeated NAMEs")
    plt.title("Repeated NAME count vs. Threshold")
    plt.grid(True)
    plt.xticks(np.arange(0,1.01,0.1))
    plt.minorticks_on()
    plt.savefig(os.path.join(output_dir, "repeated_count_vs_threshold.png"), dpi=150)
    plt.close()

    ################################################################
    # F) Combined Plot (no row=unique annotation, shorter vertical lines)
    ################################################################
    plt.figure(figsize=(8,6))
    plt.plot(summary_df["threshold"], summary_df["row_count"], marker='o', label="Total Row Count")
    plt.plot(summary_df["threshold"], summary_df["unique_count"], marker='s',color="green", label="Unique NAMEs")
    plt.plot(summary_df["threshold"], summary_df["repeated_count"], marker='^',color="red", label="Repeated NAMEs")

    # We'll define a helper that draws a vertical line from y=0 up to the max
    # among row_count, unique_count, repeated_count *for that threshold*.
    def draw_vertical_segment(thr_val, color, label_txt):
        # find the row in summary_df
        row_match = summary_df[summary_df["threshold"] == thr_val]
        if row_match.empty:
            return
        # get max among row_count, unique_count, repeated_count
        row_data = row_match.iloc[0]
        y_lim = max(row_data["row_count"], row_data["unique_count"], row_data["repeated_count"])
        # draw a partial line from y=0 to y=y_lim
        plt.vlines(x=thr_val, ymin=0, ymax=y_lim, color=color, linestyles='--', label=label_txt)

    # If both thresholds exist and are the same, unify
    if (
       not np.isnan(full_analysis_threshold) 
       and not np.isnan(core_level_threshold) 
       and abs(full_analysis_threshold - core_level_threshold) < 1e-9
    ):
        # same threshold
        draw_vertical_segment(full_analysis_threshold, "purple", f"Full+Core thr={full_analysis_threshold:.2f}")
        core_level_threshold = float("nan")
        # full_analysis_threshold = float("nan")

    # draw them if still valid
    if not np.isnan(full_analysis_threshold):
        draw_vertical_segment(full_analysis_threshold, "purple", f"Full thr={full_analysis_threshold:.2f}")
    if not np.isnan(core_level_threshold):
        draw_vertical_segment(core_level_threshold, "orange", f"Core thr={core_level_threshold:.2f}")

    plt.xlabel("Threshold")
    plt.ylabel("Count")
    plt.title("Combined Plot: Total Rows, Unique Rows, Repeated Rows vs. Threshold")
    plt.legend()
    plt.grid(True)
    plt.xticks(np.arange(0,1.01,0.1))
    plt.minorticks_on()
    plt.savefig(os.path.join(output_dir, "combined_plot.png"), dpi=150)
    plt.close()

    ################################################################
    # G) Save summary_df => CSV
    ################################################################
    summary_csv_path = os.path.join(output_dir, "thresholds_summary.csv")
    summary_df.to_csv(summary_csv_path, index=False)
    print(f"[INFO] Saved summary => {summary_csv_path}")

    ################################################################
    # H) Save summary_df => DOCX
    ################################################################
    docx_path = os.path.join(output_dir, "thresholds_summary.docx")
    doc = Document()
    doc.add_heading("Thresholds Summary (Final Logic)", 0)

    doc.add_paragraph(
        "This document summarizes the threshold-based deduplication results.\n"
        "We define:\n"
        " - Full Analysis: no repeated entries + maximum unique_count (tie → highest threshold)\n"
        " - Core-Level Analysis: no repeated entries + row_count ≤ 9 (tie → highest threshold)\n"
        "All other thresholds are 'partial analysis'.\n"
        "The combined plot no longer includes any 'Row=Unique' annotations.\n"
        "Vertical lines for Full/Core thresholds are drawn only from y=0 to the maximum of the three metrics "
        "at that threshold, rather than the entire axis."
    )

    table = doc.add_table(rows=1, cols=len(summary_df.columns))
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(summary_df.columns):
        hdr_cells[i].text = str(col_name)

    for _, row_data in summary_df.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(summary_df.columns):
            row_cells[i].text = str(row_data[col_name])

    doc.add_paragraph("\nChosen thresholds:")

    # Full Analysis
    if np.isnan(full_analysis_threshold):
        doc.add_paragraph("  - Full Analysis: NONE found (no threshold yields zero repeats?).")
    else:
        doc.add_paragraph(f"  - Full Analysis threshold = {full_analysis_threshold:.2f}.")

    # Core-Level
    if np.isnan(core_level_threshold):
        doc.add_paragraph("  - Core-Level Analysis: NONE found (no threshold yields no repeats & ≤9 rows?).")
    else:
        doc.add_paragraph(f"  - Core-Level Analysis threshold = {core_level_threshold:.2f}.")

    doc.add_paragraph(
        "All other thresholds that don't meet either condition are 'partial analysis'.\n"
        "If Full and Core turned out the same, we unify them with one line labeled 'Full+Core'."
    )

    doc.save(docx_path)
    print(f"[INFO] Saved Word doc => {docx_path}")
    print("[INFO] Done. All final logic and outputs generated.")

    ################################################################
    # I)  RETURN SECTION  (new)
    ################################################################
    # If no Full‑Analysis threshold was found, return None so callers
    # can detect the situation.
    full_thr_key = (
        f"{full_analysis_threshold:.2f}"
        if not np.isnan(full_analysis_threshold) else None
    )

    core_thr_key = (
        f"{core_level_threshold:.2f}"
        if not np.isnan(core_level_threshold) else None
    )

    full_df = thr2df.get(full_thr_key) if full_thr_key else None
    core_df = thr2df.get(core_thr_key) if core_thr_key else None

    return full_df, core_df, summary_df

    


############################################################################
# USAGE FUNCTION (if needed)
############################################################################
def run_monte_carlo_demo(your_dataframe: pd.DataFrame, dir_name: str):
    """
    Example usage: run_monte_carlo_demo(WO_2023060387_A1,"WO_2023060387_A1_4")

    This runs the new final analysis:
      - Removes "Row=Unique" annotations
      - Shortens vertical lines for Full/Core thresholds
      - Outputs .csv, .docx, and 4 plot PNGs in 'dir_name'
    """
    monte_carlo_analysis(
        df=your_dataframe,
        output_dir=dir_name,
        name_col="NAME",
        model_name=DEFAULT_EMBED_MODEL,
        threshold_step=0.05
    )


#!/usr/bin/env python
# JUPYTER NOTEBOOK CELL START

"""
Multi-Patent Product Matching Pipeline Using OpenAI Embeddings
(with Dictionary/Fuzzy fallback & user-defined CSV headings)

CHANGES:
 - We now use `split_entries` (falling back to `representative_member`) as the
   display name for each patent chemical in the summary CSV.
 - We collect synonyms from representative_member, standard_names, and split_entries,
   but the 'primary' name we show is from `split_entries` if available.
 - In the final summary, for each product we add two columns:
    1) product_{id}_matches -> "ingredient1(score), ingredient2(score), ..."
    2) product_{id}_pct     -> that product's overall match percentage
 - This ensures we see both the matched ingredients & scores per chemical and
   the overall fraction matched (the same fraction shown in the top-K CSV).

NEW MODIFICATION:
 - Added start_row_num and end_row_num to limit processing to a subset
   of the products.csv. Raises IndexError if out of range.

ADDITIONAL NEW MODIFICATION:
 - Reuse local embeddings from 'product_ingredients_embeddings.json' if present;
   if not found, call OpenAI to compute, then store it in the cache.
"""

import os
import re
import json
import openai
import numpy as np
import pandas as pd
from typing import List, Dict, Any, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed

# For fallback matching
from rapidfuzz import process as fuzzy_process

# For cosine similarity
from numpy.linalg import norm


# ---------------- CONFIG ----------------
PRODUCT_FILE = "products.csv"         # Common products file
CHUNK_SIZE = 1000
TOP_K = 300
MAX_WORKERS = 20

EMBED_SCORE_THRESHOLD = 0.75   # If below this, fallback to dictionary/fuzzy
OPENAI_EMBEDDING_MODEL = "text-embedding-3-large"  # Use your desired model

# Default placeholders
DEFAULT_ONLY_CLAIM = "Only Claim"
DEFAULT_COMPREHENSION_LEVEL = "Middle-Level"

# NEW: Path to local ingredient embeddings cache (JSON).
CACHED_EMBEDDINGS_PATH = "product_ingredients_embeddings.json"

# NEW: Load any existing embeddings cache into memory
if os.path.exists(CACHED_EMBEDDINGS_PATH):
    try:
        with open(CACHED_EMBEDDINGS_PATH, "r", encoding="utf-8") as f:
            EMBEDDINGS_CACHE = json.load(f)
        # Ensure all keys are strings, all values are list[float].
        # If you'd like to handle type-checking or potential corruption, do that here.
        print(f"[INFO] Loaded existing embeddings cache with {len(EMBEDDINGS_CACHE)} entries.")
    except Exception as e:
        print(f"[WARN] Could not read from {CACHED_EMBEDDINGS_PATH}: {e}")
        EMBEDDINGS_CACHE = {}
else:
    EMBEDDINGS_CACHE = {}


# ============== OPENAI EMBEDDING ==============
def embed_text(text: str) -> List[float]:
    """
    Calls openai.embeddings.create for a single text
    with model=OPENAI_EMBEDDING_MODEL.

    If the text is already in EMBEDDINGS_CACHE, reuse that vector
    and skip the API call. Otherwise, call OpenAI and store the result.
    """
    text_key = text.strip().lower()  # or simply text.strip()
    
    # 1. Check cache first
    if text_key in EMBEDDINGS_CACHE:
        return EMBEDDINGS_CACHE[text_key]
        print("The text_key not found ", text_key)

    # 2. Otherwise, call OpenAI Embeddings API
    try:
        print("The text being sent is:", text)
        response = openai.embeddings.create(
            model=OPENAI_EMBEDDING_MODEL,
            input=text
        )
         
        emb_vec = response.data[0].embedding

        # 3. Store in cache
        EMBEDDINGS_CACHE[text_key] = emb_vec
        return emb_vec

    except Exception as e:
        print(f"[Error] Embedding failed for text: {text[:30]} => {e}")
        return []


def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    an = norm(a)
    bn = norm(b)
    if an == 0 or bn == 0:
        return 0.0
    return float(np.dot(a, b) / (an*bn))


# =========== PARSE & EMBED PATENT CHEMICALS ===========
def parse_synonyms_from_row(row: pd.Series) -> List[str]:
    """
    Gather synonyms from:
      - representative_member
      - standard_names (JSON or list)
      - split_entries
    """
    synonyms = set()

    # representative_member
    rep = str(row.get("representative_member", "")).strip()
    if rep:
        synonyms.add(rep)

    # standard_names
    std_names = row.get("standard_names", "")
    if isinstance(std_names, str):
        try:
            maybe_list = json.loads(std_names)
            if isinstance(maybe_list, list):
                for x in maybe_list:
                    synonyms.add(str(x).strip())
            else:
                synonyms.add(std_names.strip())
        except:
            synonyms.add(std_names.strip())
    elif isinstance(std_names, list):
        for x in std_names:
            synonyms.add(str(x).strip())

    # split_entries
    sp_entries = row.get("split_entries", "")
    if isinstance(sp_entries, str):
        sp_entries = sp_entries.strip()
        if sp_entries:
            synonyms.add(sp_entries)
    elif isinstance(sp_entries, list):
        for x in sp_entries:
            synonyms.add(str(x).strip())

    final_syns = []
    for s in synonyms:
        c = s.strip()
        if c:
            final_syns.append(c)
    return list(set(final_syns))

def embed_synonyms(syn_list: List[str]) -> np.ndarray:
    """
    Embed each synonym individually and average them into one vector.
    Returns a 1536-dim vector or zero-vector if none.
    
    *Uses the embed_text() function which now consults EMBEDDINGS_CACHE first.*
    """
    if not syn_list:
        return np.zeros((1536,), dtype=float)

    embs = []
    for syn in syn_list:
        v = embed_text(syn)
        if v:
            embs.append(v)

    if not embs:
        return np.zeros((1536,), dtype=float)
    arr = np.array(embs)
    return np.mean(arr, axis=0)


# =========== FALLBACK DICTIONARY/FUZZY ===========
def dictionary_exact_match(ingredient: str, synonyms_map: Dict[str,int]) -> Tuple[int,bool]:
    low = ingredient.lower().strip()
    if low in synonyms_map:
        return (synonyms_map[low], True)
    return (None, False)

def dictionary_fuzzy_match(ingredient: str, synonyms_map: Dict[str,int], threshold=85) -> Tuple[int,bool]:
    low = ingredient.lower().strip()
    if not synonyms_map:
        return (None,False)
    all_syns = list(synonyms_map.keys())
    best_match, best_score, _ = fuzzy_process.extractOne(low, all_syns)
    if best_score >= threshold:
        return (synonyms_map[best_match], True)
    return (None,False)

def confirm_dictionary_match(ingredient: str, pat_id: int, synonyms_map: Dict[str,int]) -> bool:
    """
    If embedding matched pat_id, see if dictionary exact/fuzzy yields same ID.
    """
    did, ex = dictionary_exact_match(ingredient, synonyms_map)
    if ex and did == pat_id:
        return True
    if not ex:
        did2, fz = dictionary_fuzzy_match(ingredient, synonyms_map)
        if fz and did2 == pat_id:
            return True
    return False


# =========== PRODUCT INGREDIENT PARSER ===========
def parse_ingredient_list(ing_str: str) -> List[str]:
    if not ing_str:
        return []
    ing_str = ing_str.replace("\n"," ")
    splitted = re.split(r",|;|/|\r", ing_str)
    return [x.strip() for x in splitted if x.strip()]


# =========== MATCH ONE INGREDIENT ===========
def match_ingredient(ingredient: str,
                     patent_master: List[Dict[str,Any]],
                     synonyms_map: Dict[str,int]) -> Tuple[int,float,bool]:
    """
    1) Embed the ingredient with OpenAI (or from cache).
    2) Compare to all patent vectors => best ID & best sim.
    3) If best sim >= EMBED_SCORE_THRESHOLD => embed match
       else fallback to dictionary/fuzzy => (score=0.5).
    Returns (patent_id, score, dictionary_confirmation).
    """
    vec = np.array(embed_text(ingredient), dtype=float)
    if vec.size == 0:
        # embed call failed => dictionary fallback
        d_id, d_ex = dictionary_exact_match(ingredient, synonyms_map)
        if not d_ex:
            d_id2, d_fz = dictionary_fuzzy_match(ingredient, synonyms_map)
            if d_fz:
                d_id = d_id2
        if d_id is not None:
            return (d_id, 0.5, True)
        else:
            return (None, 0.0, False)

    # find best match among all
    best_pid = None
    best_score = 0.0
    for pm in patent_master:
        p_vec = pm["vector"]
        pid = pm["id"]
        sim = cosine_similarity(vec, p_vec)
        if sim > best_score:
            best_score = sim
            best_pid = pid

    if best_score >= EMBED_SCORE_THRESHOLD:
        # accepted embed match
        confirm = confirm_dictionary_match(ingredient, best_pid, synonyms_map)
        return (best_pid, best_score, confirm)
    else:
        # fallback
        d_id, d_ex = dictionary_exact_match(ingredient, synonyms_map)
        if not d_ex:
            d_id2, d_fz = dictionary_fuzzy_match(ingredient, synonyms_map)
            if d_fz:
                d_id = d_id2
        if d_id is not None:
            return (d_id, 0.5, True)
        else:
            return (None, 0.0, False)


# =========== PROCESS ONE PRODUCT ROW ===========
def process_one_product(row: pd.Series,
                        idx: int,
                        total: int,
                        patent_master: List[Dict[str,Any]],
                        synonyms_map: Dict[str,int]) -> Dict[str,Any]:
    """
    For each product row, parse ingredients, attempt to match
    each ingredient to the patent chemicals, count fraction matched, etc.
    """
    p_id = row.get("product_id","")
    p_name = row.get("name","")
    p_url = row.get("url","")
    ings_str = row.get("ingredients","")

    print(f"\n[Product {idx+1}/{total}] product_id='{p_id}', name='{p_name}', url='{p_url}'")

    ings = parse_ingredient_list(ings_str)
    matched_count = 0
    matched_map = []
    for ing in ings:
        pat_id, score, dict_conf = match_ingredient(ing, patent_master, synonyms_map)
        if pat_id is not None and score > 0:
            pm = patent_master[pat_id]
            matched_count += 1
            pat_name = pm["patent_chemical"]  # we use the 'split_entries' fallback name
            matched_map.append({
                "product_ingredient": ing,
                "matched_patent_chemical": pat_name,
                "matched_chemical_synonyms": pm["synonyms"],
                "score": score,
                "dictionary_confirmation": dict_conf
            })
        else:
            matched_map.append({
                "product_ingredient": ing,
                "matched_patent_chemical": None,
                "matched_chemical_synonyms": [],
                "score": 0.0,
                "dictionary_confirmation": False
            })

    fraction = float(matched_count) / len(ings) if ings else 0.0
    pct_str = f"{fraction*100:.2f}%"
    print(f"     => Matched {pct_str} of product's ingredients to patent chemicals")

    return {
        "product_id": p_id,
        "product_name": p_name,
        "product_url": p_url,
        "fraction_matched": fraction,
        "match_percentage": pct_str,
        "matched_ingredients_map": matched_map
    }


def update_top_k_candidates(all_cands: List[Dict[str,Any]],
                            new_batch: List[Dict[str,Any]],
                            top_k: int) -> List[Dict[str,Any]]:
    """
    Merge old + new results, sort by fraction_matched desc,
    keep only top_k.
    """
    combined = all_cands + new_batch
    sorted_combined = sorted(combined, key=lambda x: x["fraction_matched"], reverse=True)
    return sorted_combined[:top_k]

from typing import Union

# =========== MAIN PER-PATENT ANALYSIS ===========
def run_analysis_for_one_patent(
    patent_csv_path: Union[str, pd.DataFrame],
    product_csv_path: str,
    only_claim_str: str,
    comprehension_level_str: str,
    start_row_num: int = None,     
    end_row_num: int = None        
) -> Tuple[str, str]:
    """
    Reads one patent CSV, parse 'split_entries' as the displayed name,
    collects synonyms, does the product scan, writes top-K + summary CSVs.

    Supports start_row_num and end_row_num to limit the range
    of products used from product_csv_path.
    """
    # if not os.path.exists(patent_csv_path):
    #     raise FileNotFoundError(f"Cannot find patent CSV => {patent_csv_path}")

    df_patent_local = patent_csv_path.copy().fillna("")
    patent_csv_path_str = "<in‑memory>"

    # df_patent_local = pd.read_csv(patent_csv_path).fillna("")
    # Remove rows that have no patent_number

    df_patent_local = df_patent_local.dropna(subset=["patent_number"])
    df_patent_local = df_patent_local[df_patent_local["patent_number"].str.strip() != ""]
    if df_patent_local.empty:
        print(f"[WARN] Patent CSV has no valid patent_number rows => {patent_csv_path}")
        return

    patent_number_vals = df_patent_local["patent_number"].unique()
    if len(patent_number_vals) != 1:
        print(f"[ERROR] Patent CSV {patent_csv_path} has multiple or zero distinct patent_number values: {patent_number_vals}")
        return
    patent_num = str(patent_number_vals[0])

    print(f"[INFO] Starting analysis for patent_number={patent_num}, CSV={patent_csv_path}")

    # Build local patent master
    local_patent_master = []
    for idx, row in df_patent_local.iterrows():
        # For display, prefer split_entries if present, else fallback
        chem_display = str(row.get("split_entries","")).strip()
        if not chem_display:
            chem_display = str(row.get("representative_member","")).strip()

        synonyms = parse_synonyms_from_row(row)
        item = {
            "id": idx,
            "patent_chemical": chem_display,  # used in summary
            "synonyms": synonyms,
            "row_data": row.to_dict(),
            "vector": None
        }
        local_patent_master.append(item)

    print(f"[INFO] local_patent_master size={len(local_patent_master)} for patent_number={patent_num}")

    # Build dictionary for fallback
    local_synonyms_dict: Dict[str,int] = {}
    for pm in local_patent_master:
        i = pm["id"]
        for s in pm["synonyms"]:
            low_s = s.lower().strip()
            local_synonyms_dict[low_s] = i

    # Embed synonyms concurrently
    local_vectors = [None]*len(local_patent_master)

    def worker_embed_patent(idx: int, syns: List[str]) -> Tuple[int, np.ndarray]:
        emb = embed_synonyms(syns)
        return (idx, emb)

    futs = []
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as exe:
        for pm in local_patent_master:
            i = pm["id"]
            syns = pm["synonyms"]
            fut = exe.submit(worker_embed_patent, i, syns)
            futs.append(fut)

    for fut in as_completed(futs):
        i, emb = fut.result()
        local_vectors[i] = emb

    for i, pm in enumerate(local_patent_master):
        pm["vector"] = local_vectors[i]

    # Load products
    if not os.path.exists(product_csv_path):
        raise FileNotFoundError(f"Cannot find product CSV => {product_csv_path}")
    df_products = pd.read_csv(product_csv_path).fillna("")
    print(f"[INFO] Loaded products from '{product_csv_path}', shape={df_products.shape}")

    # =============== START/END ROW SLICING ===============
    total_rows = len(df_products)
    if start_row_num is None:
        start_row_num = 0
    if end_row_num is None:
        end_row_num = total_rows - 1

    if start_row_num < 0 or start_row_num >= total_rows:
        raise IndexError(f"start_row_num={start_row_num} is out of range [0, {total_rows-1}]")
    if end_row_num < start_row_num or end_row_num >= total_rows:
        raise IndexError(
            f"end_row_num={end_row_num} is out of range or less than start_row_num={start_row_num}"
        )

    # Slice the product DataFrame
    df_products = df_products.iloc[start_row_num : end_row_num+1].reset_index(drop=True)
    # =============== END NEW CODE ===============

    num_products = len(df_products)
    print(f"[INFO] After slicing, we have {num_products} products to process.")

    # Process in chunks
    top_candidates: List[Dict[str,Any]] = []
    start_idx = 0

    while start_idx < num_products:
        chunk_end = min(start_idx + CHUNK_SIZE, num_products)
        df_chunk = df_products.iloc[start_idx:chunk_end].reset_index(drop=True)
        print(f"[INFO] Processing chunk from product index {start_idx} to {chunk_end}...")

        chunk_results = []
        futs2 = []
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as exe:
            for i, row in df_chunk.iterrows():
                fut = exe.submit(
                    process_one_product, 
                    row, 
                    i, 
                    len(df_chunk), 
                    local_patent_master, 
                    local_synonyms_dict
                )
                futs2.append(fut)

        done_count = 0
        for fut in as_completed(futs2):
            try:
                res = fut.result()
                chunk_results.append(res)
            except Exception as e:
                print(f"[ERROR] pipeline error: {e}")
            done_count += 1
        print(f"   -> Completed {done_count}/{len(df_chunk)} in this chunk.")

        top_candidates = update_top_k_candidates(top_candidates, chunk_results, TOP_K)
        start_idx = chunk_end
        print(f"[INFO] After scanning up to {chunk_end}, top_candidates size={len(top_candidates)}\n")

    final_topk = sorted(top_candidates, key=lambda x: x["fraction_matched"], reverse=True)
    if not final_topk:
        print(f"[INFO] No matched products for patent_number={patent_num}. Exiting.")
        return

    # Write products_topk_matched_{patent_num}.csv
    matched_csv_filename = f"products_topk_matched_{patent_num}.csv"
    print(f"[INFO] Building final TOP-K DataFrame => {matched_csv_filename}")

    rows_for_csv = []
    for r in final_topk:
        maps_json = json.dumps(r["matched_ingredients_map"], ensure_ascii=False)
        rows_for_csv.append({
            "patent_number": patent_num,
            "product_id": r["product_id"],
            "product_name": r["product_name"],
            "product_url": r["product_url"],
            "match_percentage": r["match_percentage"],
            "matched_ingredients_map": maps_json
        })
    df_final_topk = pd.DataFrame(rows_for_csv, columns=[
        "patent_number","product_id","product_name","product_url","match_percentage","matched_ingredients_map"
    ])

    # heading_line = (
    #     f"# Patent Number: {patent_num}, Only Claim: {only_claim_str}, "
    #     f"Comprehension Level: {comprehension_level_str}\n"
    # )
    # with open(matched_csv_filename, "w", encoding="utf-8") as f:
    #     f.write(heading_line)
    df_final_topk.to_csv(matched_csv_filename, index=False, encoding="utf-8")
    print(f"[INFO] Saved top-K CSV => '{matched_csv_filename}'")

    # Write products_topk_summary_{patent_num}.csv
    summary_csv_filename = f"products_topk_summary_{patent_num}.csv"
    print(f"[INFO] Building summary => {summary_csv_filename}")

    summary_rows = []
    for pm in local_patent_master:
        # We'll use pm["patent_chemical"] (from split_entries if available)
        pm_name = pm["patent_chemical"]
        rowd = {
            "patent_id": pm["id"],
            "patent_number": patent_num,
            "patent_chemical": pm_name
        }
        # For each product in final_topk, we store 2 columns:
        # 1) product_{id}_matches: matched ingredients (with scores) for this chemical
        # 2) product_{id}_pct: that product's overall matched fraction
        for r in final_topk:
            pcol_matches = f"product_{r['product_id']}_matches"
            pcol_pct = f"product_{r['product_id']}_pct"
            matched_ings_with_scores = []
            for mi in r["matched_ingredients_map"]:
                if mi["matched_patent_chemical"] == pm_name:
                    matched_ings_with_scores.append(f"{mi['product_ingredient']}({mi['score']:.2f})")
            rowd[pcol_matches] = ", ".join(matched_ings_with_scores) if matched_ings_with_scores else ""
            rowd[pcol_pct] = r["match_percentage"]  # overall fraction matched, same as topK

        summary_rows.append(rowd)

    df_summary = pd.DataFrame(summary_rows)

    # heading_line2 = (
    #     f"# Patent Number: {patent_num}, Only Claim: {only_claim_str}, "
    #     f"Comprehension Level: {comprehension_level_str}\n"
    # )
    # with open(summary_csv_filename, "w", encoding="utf-8") as f:
    #     f.write(heading_line2)
    df_summary.to_csv(summary_csv_filename, index=False, encoding="utf-8")
    print(f"[INFO] Saved summary CSV => '{summary_csv_filename}'")
      # --- return the two files so callers (or FastAPI background task)
    #     can hand them to the frontend for download -------------------
    return os.path.abspath(matched_csv_filename), os.path.abspath(summary_csv_filename)


# PATENT_MAP = {
#     "chosen_A.csv": "US1091278B2",
#     "chosen_B.csv": "US2009274638A1",
#     "chosen_C.csv": "US2010168049A1",
#     "chosen_D.csv": "US2010168055A1",
#     "chosen_E.csv": "WO2023060387A1",
# }

def ensure_patent_number_in_csv(csv_path: Union[str, pd.DataFrame], patent_number: str):
    """
    Check if the CSV file at 'csv_path' contains a valid 'patent_number' column.
    If not, or if the column has empty values, fill it with the pre-defined
    patent number from PATENT_MAP based on the CSV filename.
    """
    if isinstance(csv_path, pd.DataFrame):
        # column already exists? – nothing to do
        if 'patent_number' not in csv_path.columns or csv_path['patent_number'].isna().any():
            csv_path['patent_number'] = patent_number
        return
    
    df = pd.read_csv(csv_path)

    # Extract just the filename from the path (e.g. "chosen_A.csv").
    filename = os.path.basename(csv_path)

    # Determine the patent number to use from the map; if missing, default to a placeholder.
    patent_number = patent_number

    # If the patent_number column doesn't exist or is empty/placeholder, fill it in.
    if 'patent_number' not in df.columns:
        print(f"[WARNING] '{csv_path}' does not have a 'patent_number' column. Adding one with '{patent_number}'.")
        df['patent_number'] = patent_number
        df.to_csv(csv_path, index=False)
        print(f"[INFO] Updated '{csv_path}' with patent_number='{patent_number}'")

    else:
        # Check if any rows have empty or placeholder values.
        empty_or_placeholder = (
            df['patent_number'].isnull() |
            (df['patent_number'].astype(str).str.strip() == "") |
            (df['patent_number'] == "INSERT_PATENT_NUMBER_HERE")
        )
        if empty_or_placeholder.any():
            print(f"[WARNING] '{csv_path}' has empty or placeholder 'patent_number' values. "
                  f"Replacing them with '{patent_number}'.")
            df.loc[empty_or_placeholder, 'patent_number'] = patent_number
            df.to_csv(csv_path, index=False)
            print(f"[INFO] Updated '{csv_path}' with patent_number='{patent_number}'")


def run_analysis_for_multiple_csvs(
    patent_csv_paths: Union[pd.DataFrame, List[Union[str, pd.DataFrame]]],
    product_csv_path: str = PRODUCT_FILE,
    default_only_claim: str = DEFAULT_ONLY_CLAIM,
    default_comprehension_level: str = DEFAULT_COMPREHENSION_LEVEL,
    start_row_num: int = None,
    end_row_num: int = None
) -> List[Tuple[str,str]]:
    """
    Handle multiple patent CSVs, producing two output CSVs per patent:
      products_topk_matched_{patent_number}.csv
      products_topk_summary_{patent_number}.csv
    Each with headings referencing 'Only Claim' and 'Comprehension Level'.

    Now supports start_row_num and end_row_num to limit the rows used
    from product_csv_path for each patent analysis.
    """
    all_reports: List[Tuple[str,str]] = []
    if isinstance(patent_csv_paths, (pd.DataFrame, str)):
        patent_csv_paths = [patent_csv_paths] 

    for csv_p in patent_csv_paths:
        # Ensure the CSV has a valid patent_number column (hard-coded from PATENT_MAP if needed).
        ensure_patent_number_in_csv(csv_p, "number_1")

    for csv_p in patent_csv_paths:
        reports = run_analysis_for_one_patent(
            patent_csv_path=csv_p,
            product_csv_path=product_csv_path,
            only_claim_str=default_only_claim,
            comprehension_level_str=default_comprehension_level,
            start_row_num=start_row_num,
            end_row_num=end_row_num
        )
        print(f"[INFO] Completed patent analysis for => {csv_p}")
        all_reports.append(reports) 

    return all_reports

###############################################################
# The End of the Pipeline code
###############################################################


